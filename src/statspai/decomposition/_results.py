"""
Shared mixin for decomposition result classes.

Every result class in :mod:`statspai.decomposition` benefits from a small
common surface area: structured exports (Excel, Word, dict, JSON), a
canonical bibliography lookup (:meth:`cite`), normal-approximation
confidence intervals from a stored ``se``-dict, and a colour palette
shared with :mod:`.plots` so that every figure looks like part of the
same family.

The mixin is intentionally permissive — a result class only needs to
expose a subset of canonical attributes and the relevant exporter will
try its best. Where the necessary information is missing we degrade to
"not available" rather than raising.

Result classes opt in by inheriting from :class:`DecompResultMixin` (or
adding it as an additional base for ``@dataclass`` results) and, ideally,
by populating the canonical attributes:

- ``overall`` *(dict)* — at minimum ``gap``, plus any of ``explained``,
  ``unexplained``, ``composition``, ``structure`` (and matching ``_se``
  entries for analytical CIs).
- ``detailed`` *(pandas.DataFrame)* — variable-level contributions, with
  a numeric ``se`` column where available.
- ``method_name`` *(str)* — pretty label for headings.
- ``bib_keys`` *(list[str])* — citation keys from ``paper.bib``.

If only ``overall`` (or only ``detailed``) is present, the exporter
silently omits the missing pane. The mixin never overrides an existing
``summary``, ``plot``, ``to_latex``, or ``_repr_html_`` defined on the
subclass; it only fills in the gaps.
"""
from __future__ import annotations

import contextlib
from dataclasses import is_dataclass, asdict
from typing import Any, Dict, Iterable, Mapping, Optional, Tuple, Union
import io
import json
import warnings

import numpy as np
import pandas as pd
from scipy import stats


# ════════════════════════════════════════════════════════════════════════
# Shared visual style for every decomposition plot
# ════════════════════════════════════════════════════════════════════════

#: Canonical palette used across decomposition plots so that every figure
#: looks like part of the same family. ``pos`` / ``neg`` colour-code the
#: sign of a contribution; ``a`` / ``b`` mark the two groups; ``cf`` is
#: the counterfactual; ``accent`` is for total / gap bars.
DECOMP_PALETTE: Dict[str, str] = {
    "pos": "#1976D2",       # contribution > 0 (Material blue 700)
    "neg": "#E53935",       # contribution < 0 (Material red 600)
    "a": "#1E88E5",         # group A
    "b": "#FB8C00",         # group B
    "cf": "#43A047",        # counterfactual
    "accent": "#37474F",    # total / gap (Material blue-grey 800)
    "ci": "#90A4AE",        # confidence-interval whiskers
    "between": "#1976D2",
    "within": "#E53935",
    "overlap": "#8E24AA",
}


def apply_decomp_style(ax) -> None:
    """In-place styling: light grid, grey spines, despined top/right.

    Idempotent — safe to call repeatedly. Used by every plot helper to
    give the module a consistent look without forcing a global
    matplotlib style change on the user.
    """
    ax.grid(True, axis="x", color="#ECEFF1", linewidth=0.6, zorder=0)
    ax.grid(True, axis="y", color="#ECEFF1", linewidth=0.6, zorder=0)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    for s in ("left", "bottom"):
        ax.spines[s].set_color("#90A4AE")
        ax.spines[s].set_linewidth(0.8)
    ax.tick_params(colors="#37474F", which="both", length=3)


# ════════════════════════════════════════════════════════════════════════
# Bibliography lookup
# ════════════════════════════════════════════════════════════════════════

#: BibTeX-style citations for every method in ``sp.decomposition``. Keys
#: line up with entries in ``paper.bib`` (verified via Crossref / DOI /
#: arXiv) so that ``result.cite()`` returns the same canonical strings
#: that downstream documentation uses.
_CITATIONS: Dict[str, str] = {
    # Mean decomposition ----------------------------------------------
    "blinder1973wage": (
        "Blinder, A. S. (1973). Wage Discrimination: Reduced Form and "
        "Structural Estimates. Journal of Human Resources, 8(4), 436-455. "
        "doi:10.2307/144855"
    ),
    "oaxaca1973male": (
        "Oaxaca, R. (1973). Male-Female Wage Differentials in Urban Labor "
        "Markets. International Economic Review, 14(3), 693-709. "
        "doi:10.2307/2525981"
    ),
    "neumark1988employers": (
        "Neumark, D. (1988). Employers' Discriminatory Behavior and the "
        "Estimation of Wage Discrimination. Journal of Human Resources, "
        "23(3), 279-295. doi:10.2307/145830"
    ),
    "cotton1988estimation": (
        "Cotton, J. (1988). On the Decomposition of Wage Differentials. "
        "Review of Economics and Statistics, 70(2), 236-243. "
        "doi:10.2307/1928307"
    ),
    "reimers1983labor": (
        "Reimers, C. W. (1983). Labor Market Discrimination Against "
        "Hispanic and Black Men. Review of Economics and Statistics, "
        "65(4), 570-579. doi:10.2307/1935925"
    ),
    "jann2008blinder": (
        "Jann, B. (2008). The Blinder-Oaxaca Decomposition for Linear "
        "Regression Models. Stata Journal, 8(4), 453-479. "
        "doi:10.1177/1536867X0800800401"
    ),
    "gelbach2016covariates": (
        "Gelbach, J. B. (2016). When Do Covariates Matter? And Which "
        "Ones, and How Much? Journal of Labor Economics, 34(2), 509-543. "
        "doi:10.1086/683668"
    ),
    "fairlie2005extension": (
        "Fairlie, R. W. (2005). An Extension of the Blinder-Oaxaca "
        "Decomposition Technique to Logit and Probit Models. Journal of "
        "Economic and Social Measurement, 30(4), 305-316. "
        "doi:10.3233/JEM-2005-0259"
    ),
    "yun2004decomposing": (
        "Yun, M.-S. (2004). Decomposing differences in the first moment. "
        "Economics Letters, 82(2), 275-280. "
        "doi:10.1016/j.econlet.2003.09.008"
    ),
    "bauer2008extension": (
        "Bauer, T. K. & Sinning, M. (2008). An extension of the "
        "Blinder-Oaxaca decomposition to nonlinear models. AStA Advances "
        "in Statistical Analysis, 92(2), 197-206. "
        "doi:10.1007/s10182-008-0056-3"
    ),
    "kline2011oaxaca": (
        "Kline, P. (2011). Oaxaca-Blinder as a Reweighting Estimator. "
        "American Economic Review: Papers & Proceedings, 101(3), 532-537. "
        "doi:10.1257/aer.101.3.532"
    ),
    # Distributional --------------------------------------------------
    "firpo2009unconditional": (
        "Firpo, S., Fortin, N. M., & Lemieux, T. (2009). Unconditional "
        "Quantile Regressions. Econometrica, 77(3), 953-973. "
        "doi:10.3982/ECTA6822"
    ),
    "firpo2018decomposing": (
        "Firpo, S., Fortin, N. M., & Lemieux, T. (2018). Decomposing "
        "Wage Distributions Using Recentered Influence Function "
        "Regressions. Econometrics, 6(2), 28. "
        "doi:10.3390/econometrics6020028"
    ),
    "dinardo1996labor": (
        "DiNardo, J., Fortin, N. M., & Lemieux, T. (1996). Labor Market "
        "Institutions and the Distribution of Wages, 1973-1992: A "
        "Semiparametric Approach. Econometrica, 64(5), 1001-1044. "
        "doi:10.2307/2171954"
    ),
    "machado2005counterfactual": (
        "Machado, J. A. F. & Mata, J. (2005). Counterfactual decomposition "
        "of changes in wage distributions using quantile regression. "
        "Journal of Applied Econometrics, 20(4), 445-465. "
        "doi:10.1002/jae.788"
    ),
    "melly2005decomposition": (
        "Melly, B. (2005). Decomposition of differences in distribution "
        "using quantile regression. Labour Economics, 12(4), 577-590. "
        "doi:10.1016/j.labeco.2005.05.006"
    ),
    "chernozhukov2013inference": (
        "Chernozhukov, V., Fernandez-Val, I., & Melly, B. (2013). "
        "Inference on Counterfactual Distributions. Econometrica, 81(6), "
        "2205-2268. doi:10.3982/ECTA10582"
    ),
    "fortin2011decomposition": (
        "Fortin, N., Lemieux, T., & Firpo, S. (2011). Decomposition "
        "Methods in Economics. Handbook of Labor Economics, 4A, 1-102. "
        "doi:10.1016/S0169-7218(11)00407-2"
    ),
    "riosavila2020rif": (
        "Rios-Avila, F. (2020). Recentered influence functions (RIFs) in "
        "Stata: RIF regression and RIF decomposition. Stata Journal, "
        "20(1), 51-94. doi:10.1177/1536867X20909690"
    ),
    "cowell2007income": (
        "Cowell, F. A. & Flachaire, E. (2007). Income distribution and "
        "inequality measurement: The problem of extreme values. Journal "
        "of Econometrics, 141(2), 1044-1072. "
        "doi:10.1016/j.jeconom.2007.01.001"
    ),
    "essama2012influence": (
        "Essama-Nssah, B. & Lambert, P. J. (2012). Influence functions "
        "for policy impact analysis. In Bishop, J. A. & Salas, R. (eds), "
        "Inequality, Mobility and Segregation: Essays in Honor of Jacques "
        "Silber (Research on Economic Inequality, Vol. 20). "
        "doi:10.1108/S1049-2585(2012)0000020009"
    ),
    # Inequality ------------------------------------------------------
    "shorrocks2013decomposition": (
        "Shorrocks, A. F. (2013). Decomposition procedures for "
        "distributional analysis: A unified framework based on the "
        "Shapley value. Journal of Economic Inequality, 11(1), 99-126. "
        "doi:10.1007/s10888-011-9214-z"
    ),
    "lerman1985income": (
        "Lerman, R. I. & Yitzhaki, S. (1985). Income Inequality Effects "
        "by Income Source: A New Approach and Applications to the United "
        "States. Review of Economics and Statistics, 67(1), 151-156. "
        "doi:10.2307/1928447"
    ),
    "shorrocks1980class": (
        "Shorrocks, A. F. (1980). The Class of Additively Decomposable "
        "Inequality Measures. Econometrica, 48(3), 613-625. "
        "doi:10.2307/1913126"
    ),
    # Demographic / standardisation ----------------------------------
    "kitagawa1955components": (
        "Kitagawa, E. M. (1955). Components of a Difference Between Two "
        "Rates. Journal of the American Statistical Association, 50(272), "
        "1168-1194. doi:10.1080/01621459.1955.10501299"
    ),
    "dasgupta1993standardization": (
        "Das Gupta, P. (1993). Standardization and Decomposition of "
        "Rates: A User's Manual. U.S. Bureau of the Census, "
        "P23-186. https://www.census.gov/library/publications/1993/demo/p23-186.html"
    ),
    "kroger2021kitagawa": (
        "Kroger, H. & Hartmann, J. (2021). Extending the "
        "Kitagawa-Oaxaca-Blinder decomposition approach to panel data. "
        "Stata Journal, 21(2), 360-410. "
        "doi:10.1177/1536867X211025800"
    ),
    "oaxaca2025meets": (
        "Oaxaca, R. L. & Sierminska, E. (2025). Oaxaca-Blinder meets "
        "Kitagawa: What is the link? PLOS ONE, 20(5), e0321874. "
        "doi:10.1371/journal.pone.0321874"
    ),
    # Causal decomposition -------------------------------------------
    "lundberg2021gap": (
        "Lundberg, I. (2022). The Gap-Closing Estimand: A Causal "
        "Approach to Study Interventions That Close Disparities Across "
        "Social Categories. Sociological Methods & Research, 53(2), "
        "507-570. doi:10.1177/00491241211055769"
    ),
    "vanderweele2014unification": (
        "VanderWeele, T. J. (2014). A Unification of Mediation and "
        "Interaction: A 4-Way Decomposition. Epidemiology, 25(5), "
        "749-761. doi:10.1097/EDE.0000000000000121"
    ),
    "jackson2018decomposition": (
        "Jackson, J. W. & VanderWeele, T. J. (2018). Decomposition "
        "Analysis to Identify Intervention Targets for Reducing "
        "Disparities. Epidemiology, 29(6), 825-835. "
        "doi:10.1097/EDE.0000000000000901"
    ),
    "yu2025nonparametric": (
        "Yu, A. & Elwert, F. (2025). Nonparametric causal decomposition "
        "of group disparities. Annals of Applied Statistics, 19(1), "
        "821-845. doi:10.1214/24-AOAS1990"
    ),
    "park2024choosing": (
        "Park, S., Kang, S., & Lee, C. (2024). Choosing an Optimal "
        "Method for Causal Decomposition Analysis with Continuous "
        "Outcomes: A Review and Simulation Study. Sociological Methods "
        "& Research, 53(2), 571-615. "
        "doi:10.1177/00811750231183711"
    ),
    "ahrens2025model": (
        "Ahrens, A., Hansen, C. B., Schaffer, M. E., & Wiemann, T. "
        "(2025). Model averaging and double machine learning. Journal "
        "of Applied Econometrics, 40(3), 249-269. "
        "doi:10.1002/jae.3103"
    ),
}


# ════════════════════════════════════════════════════════════════════════
# Mixin
# ════════════════════════════════════════════════════════════════════════

class DecompResultMixin:
    """Common surface for every decomposition result class.

    Adds ``confint``, ``cite``, ``to_dict``, ``to_json``, ``to_excel``,
    ``to_word`` while leaving any existing ``summary`` / ``to_latex`` /
    ``plot`` / ``_repr_html_`` defined on the subclass untouched.
    """

    #: Pretty label used in plot titles / Word headings.
    method_name: str = "Decomposition"
    #: BibTeX keys (must match :data:`_CITATIONS` and ``paper.bib``).
    bib_keys: Tuple[str, ...] = ()

    # ── Confidence intervals from analytical SE ─────────────────────

    def confint(
        self,
        alpha: float = 0.05,
        which: str = "overall",
    ) -> Optional[Dict[str, Tuple[float, float]]]:
        """Normal-approx CIs from ``self.overall`` ``*_se`` entries.

        Returns ``None`` if the result lacks SEs; downstream code can
        fall back on bootstrap intervals stored elsewhere.
        """
        z = float(stats.norm.ppf(1 - alpha / 2))
        if which == "overall":
            o = getattr(self, "overall", None)
            if not isinstance(o, Mapping):
                return None
            ci: Dict[str, Tuple[float, float]] = {}
            for key, val in o.items():
                if key.endswith("_se") or not isinstance(val, (int, float)):
                    continue
                se = o.get(f"{key}_se")
                if isinstance(se, (int, float)) and se > 0:
                    ci[key] = (val - z * se, val + z * se)
            return ci or None
        if which == "detailed":
            df = getattr(self, "detailed", None)
            if not isinstance(df, pd.DataFrame) or df.empty:
                return None
            if "se" not in df.columns:
                return None
            value_col = next(
                (c for c in ("contribution", "delta", "composition",
                             "structure", "value")
                 if c in df.columns),
                None,
            )
            if value_col is None:
                return None
            label_col = "variable" if "variable" in df.columns else df.columns[0]
            out: Dict[str, Tuple[float, float]] = {}
            for _, row in df.iterrows():
                v = float(row[value_col])
                s = float(row["se"]) if pd.notna(row["se"]) else 0.0
                out[str(row[label_col])] = (v - z * s, v + z * s)
            return out
        raise ValueError(f"unknown which={which!r}; use 'overall' or 'detailed'")

    # ── Bibliography ────────────────────────────────────────────────

    def cite(self, fmt: str = "string") -> Union[str, list]:
        """Return canonical citations for the underlying method.

        Parameters
        ----------
        fmt : {"string", "bibtex_keys", "list"}
            ``string`` joins a numbered list (default); ``bibtex_keys``
            returns the ``paper.bib`` keys; ``list`` returns the raw
            citation strings.
        """
        keys = list(getattr(self, "bib_keys", ()))
        if not keys:
            return "" if fmt == "string" else []
        if fmt == "bibtex_keys":
            return keys
        cites = [_CITATIONS.get(k, k) for k in keys]
        if fmt == "list":
            return cites
        return "\n".join(f"[{i + 1}] {c}" for i, c in enumerate(cites))

    # ── JSON-serialisable dict ──────────────────────────────────────

    def to_dict(self) -> Dict[str, Any]:
        """JSON-serialisable snapshot of the result.

        DataFrames become a list of records, numpy arrays become lists,
        and nested mappings are recursed. Non-trivial objects fall
        through ``str()``.
        """
        if is_dataclass(self):
            base = asdict(self)
        else:
            base = {
                k: v for k, v in self.__dict__.items()
                if not k.startswith("_")
            }
        return _coerce_for_json(base)

    def to_json(self, **kwargs) -> str:
        """JSON string of :meth:`to_dict` (passes ``**kwargs`` to json.dumps)."""
        kwargs.setdefault("indent", 2)
        kwargs.setdefault("default", str)
        return json.dumps(self.to_dict(), **kwargs)

    # ── Excel export ────────────────────────────────────────────────

    def to_excel(
        self,
        path: Optional[str] = None,
        *,
        engine: Optional[str] = None,
    ) -> Optional[bytes]:
        """Write a multi-sheet workbook of every panel of the result.

        ``Overall`` (named scalars), ``Detailed`` (per-variable table),
        and any other DataFrame attribute (e.g. ``quantile_grid``,
        ``cdf_grid``, ``group_stats``) become individual sheets.

        Parameters
        ----------
        path : str or None
            File path to write. If ``None``, returns the workbook bytes
            so the caller can stream them somewhere else (e.g. a Flask
            response or in-memory upload).
        engine : str or None
            Forwarded to ``pandas.ExcelWriter``; defaults to whichever
            engine pandas auto-detects.
        """
        sheets = self._dataframe_panels()
        if not sheets:
            raise RuntimeError(
                "Result has no exportable panels; nothing to write."
            )
        buf: io.BytesIO
        if path is None:
            buf = io.BytesIO()
            target: Any = buf
        else:
            target = path
        try:
            with pd.ExcelWriter(target, engine=engine) as writer:
                for name, df in sheets.items():
                    safe = name[:31]  # Excel sheet name limit
                    df.to_excel(writer, sheet_name=safe, index=False)
        except (ImportError, ValueError) as err:  # pragma: no cover
            raise ImportError(
                "to_excel requires openpyxl or xlsxwriter. Install via "
                "pip install openpyxl"
            ) from err
        if path is None:
            return buf.getvalue()
        return None

    # ── Word export (python-docx) ───────────────────────────────────

    def to_word(self, path: str) -> str:
        """Write a Word ``.docx`` summary (requires ``python-docx``).

        Returns the absolute path to the written file. Falls back to a
        clear ImportError if ``python-docx`` is not installed.
        """
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as err:  # pragma: no cover
            raise ImportError(
                "to_word requires python-docx. "
                "Install via: pip install python-docx"
            ) from err

        doc = Document()
        doc.add_heading(getattr(self, "method_name", "Decomposition"), level=1)
        # Per-class formatted summary if available. Suppress stdout so
        # that `to_word()` doesn't print the whole summary block to the
        # console as a side effect (most concrete summary() methods
        # also call print(text) before returning).
        if hasattr(self, "summary") and callable(self.summary):
            try:
                with contextlib.redirect_stdout(io.StringIO()):
                    txt = self.summary()
            except Exception:  # noqa: BLE001
                txt = repr(self)
            for line in str(txt).splitlines():
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.name = "Menlo"
                    run.font.size = Pt(9)
        # Each DataFrame panel becomes its own table.
        for name, df in self._dataframe_panels().items():
            doc.add_heading(name, level=2)
            if df.empty:
                doc.add_paragraph("(empty)")
                continue
            tbl = doc.add_table(rows=1, cols=len(df.columns))
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr[i].text = str(col)
            for _, row in df.iterrows():
                cells = tbl.add_row().cells
                for i, col in enumerate(df.columns):
                    val = row[col]
                    cells[i].text = (
                        f"{val:.4f}" if isinstance(val, (int, float))
                        and not isinstance(val, bool)
                        else str(val)
                    )
        cite_str = self.cite("string") if hasattr(self, "cite") else ""
        if cite_str:
            doc.add_heading("References", level=2)
            doc.add_paragraph(cite_str)
        doc.save(path)
        return path

    # ── Helper: collect every DataFrame-shaped panel on the instance.

    def _dataframe_panels(self) -> Dict[str, pd.DataFrame]:
        """Return the named DataFrame panels worth exporting.

        Recognises the canonical ``overall`` dict, the ``detailed`` table
        plus any other DataFrame attribute on the instance. Subclasses
        can override to customise the shape or order.
        """
        panels: Dict[str, pd.DataFrame] = {}
        overall = getattr(self, "overall", None)
        if isinstance(overall, Mapping):
            rows = []
            for k, v in overall.items():
                if isinstance(v, (int, float, np.floating, np.integer)):
                    rows.append({"quantity": k, "value": float(v)})
            if rows:
                panels["Overall"] = pd.DataFrame(rows)
        # Common dataclass scalars (gap / composition / structure / ...).
        scalar_keys = (
            "observed_gap", "counterfactual_gap", "closed_gap",
            "gap", "composition", "structure", "explained", "unexplained",
            "between", "within", "overlap", "total_change",
            "total", "nde", "nie", "cde", "total_effect",
            "controlled_direct", "reference_interaction",
            "mediated_interaction", "propn_mediated",
            "total_disparity", "initial_disparity",
            "mediator_attributable", "propn_mediator",
            "disparity", "baseline", "prevalence", "effect", "selection",
        )
        scalar_rows = []
        for k in scalar_keys:
            v = getattr(self, k, None)
            if isinstance(v, (int, float, np.floating, np.integer)):
                scalar_rows.append({"quantity": k, "value": float(v)})
        if scalar_rows and "Overall" not in panels:
            panels["Overall"] = pd.DataFrame(scalar_rows)
        # Per-method DataFrame attributes.
        candidate_attrs = (
            "detailed", "decomposition", "detailed_composition",
            "detailed_structure", "quantile_grid", "cdf_grid",
            "components", "subgroups", "table",
        )
        for attr in candidate_attrs:
            df = getattr(self, attr, None)
            if isinstance(df, pd.DataFrame) and not df.empty:
                panels[attr.replace("_", " ").title()] = df.reset_index(drop=True)
        return panels


# ════════════════════════════════════════════════════════════════════════
# Helpers
# ════════════════════════════════════════════════════════════════════════

def _coerce_for_json(obj: Any) -> Any:
    """Recursively coerce numpy / pandas types into JSON-friendly forms."""
    if isinstance(obj, dict):
        return {str(k): _coerce_for_json(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_coerce_for_json(v) for v in obj]
    if isinstance(obj, pd.DataFrame):
        return obj.replace({np.nan: None}).to_dict(orient="records")
    if isinstance(obj, pd.Series):
        return obj.replace({np.nan: None}).to_list()
    if isinstance(obj, np.ndarray):
        return obj.tolist()
    if isinstance(obj, (np.floating, np.integer, np.bool_)):
        return obj.item()
    if isinstance(obj, float) and (np.isnan(obj) or np.isinf(obj)):
        return None
    return obj
