"""``sp.paper(data, question)`` — end-to-end "data → publication draft" pipeline.

The ``causal_workflow.CausalWorkflow`` object already strings together
``diagnose → recommend → estimate → robustness``. This module is the
agent-native top layer: take a DataFrame plus a *natural-language
question* and emit a near-publishable draft (markdown / LaTeX / Word) in
a single call.

Pipeline:

    1. Parse ``question`` (lightweight regex / token heuristic) to fill
       in any missing y/treatment/instrument/cutoff hints. Explicit
       arguments always win.
    2. Run :class:`CausalWorkflow` (diagnose → recommend → estimate →
       robustness).
    3. Build a structured EDA section from :func:`sp.sumstats` (or a
       graceful inline fallback).
    4. Render to the chosen format. Markdown leverages
       :meth:`CausalWorkflow._render_markdown`; LaTeX wraps the same
       content with ``\\section{}`` boilerplate; Word delegates to the
       built-in ``to_docx`` path on the result object.

Notes on design
---------------
- This is **orchestration only**. No numerical primitives are
  re-implemented here.
- The question parser is intentionally simple — it provides hints, never
  overrides the user's explicit kwargs. Agents can also fully bypass the
  parser by passing all relevant column args directly.
- LLM calls are not made by default. The pipeline only triggers an LLM
  oracle if the user passes ``dag=`` from a prior
  :func:`sp.llm_dag_constrained` run.
"""
from __future__ import annotations

import datetime as _dt
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd

from ..output._lineage import format_provenance, get_provenance
from ._degradation import WorkflowDegradedWarning, record_degradation


__all__ = [
    "paper",
    "PaperDraft",
    "parse_question",
    "paper_from_question",
    "WorkflowDegradedWarning",
]


# --------------------------------------------------------------------- #
#  Question parser (heuristic)
# --------------------------------------------------------------------- #

# Words that strongly hint at a design when present near a column name.
_DESIGN_HINTS = {
    "did": "did",
    "difference-in-differences": "did",
    "diff-in-diff": "did",
    "parallel trends": "did",
    "rd": "rd",
    "discontinuity": "rd",
    "regression discontinuity": "rd",
    "iv": "iv",
    "instrument": "iv",
    "instrumental variable": "iv",
    "rct": "rct",
    "randomi": "rct",  # randomi(z|s)ation
    "experiment": "rct",
    "synthetic control": "synth",
    "event study": "did",
}


def parse_question(question: str, columns: List[str]) -> Dict[str, Any]:
    """Heuristic parse of a natural-language causal question.

    Returns a dict of *hints* the caller can fall back on when explicit
    column kwargs aren't provided. Never overrides explicit args.

    Parameters
    ----------
    question : str
        Natural-language question, e.g. ``"effect of training on wages"``.
    columns : list of str
        Columns of the dataset; the parser only proposes column names
        present in this list.

    Returns
    -------
    dict
        Possible keys: ``y``, ``treatment``, ``design``,
        ``instrument``, ``running_var``, ``cutoff``,
        ``raw_question``.
    """
    out: Dict[str, Any] = {"raw_question": question}
    if not isinstance(question, str) or not question.strip():
        return out
    q = question.lower()
    cols_lower = {c.lower(): c for c in columns}

    # Design hint
    for key, design in _DESIGN_HINTS.items():
        if key in q:
            out["design"] = design
            break

    # Pattern: "effect of X on Y" / "impact of X on Y" / "X on Y"
    m = re.search(
        r"(?:effect|impact|causal effect|relationship|influence)\s+"
        r"of\s+([a-z0-9_]+)\s+on\s+([a-z0-9_]+)",
        q,
    )
    if m:
        treat, y = m.group(1), m.group(2)
        if treat in cols_lower:
            out.setdefault("treatment", cols_lower[treat])
        if y in cols_lower:
            out.setdefault("y", cols_lower[y])

    # Pattern: "Y ~ X" / "Y = X"
    m = re.search(r"([a-z0-9_]+)\s*[~=]\s*([a-z0-9_]+)", q)
    if m:
        y, treat = m.group(1), m.group(2)
        if y in cols_lower:
            out.setdefault("y", cols_lower[y])
        if treat in cols_lower:
            out.setdefault("treatment", cols_lower[treat])

    # Pattern: "instrument <Z>" / "using <Z> as an instrument"
    m = re.search(
        r"(?:instrument(?:ing)?|using)\s+([a-z0-9_]+)\s+as\s+(?:an?\s+)?"
        r"instrument",
        q,
    )
    if m and m.group(1) in cols_lower:
        out["instrument"] = cols_lower[m.group(1)]
        out["design"] = "iv"

    # Pattern: "discontinuity at <c>" / "threshold <c>"
    m = re.search(
        r"(?:discontinuity|threshold|cutoff)\s+(?:at\s+)?(-?\d+\.?\d*)", q
    )
    if m:
        try:
            out["cutoff"] = float(m.group(1))
            out["design"] = "rd"
        except ValueError:
            pass
    # Pattern: "running variable <X>"
    m = re.search(r"running\s+variable\s+([a-z0-9_]+)", q)
    if m and m.group(1) in cols_lower:
        out["running_var"] = cols_lower[m.group(1)]
        out["design"] = "rd"

    return out


# --------------------------------------------------------------------- #
#  PaperDraft
# --------------------------------------------------------------------- #


@dataclass
class PaperDraft:
    """Draft causal-analysis report assembled by :func:`sp.paper`.

    Attributes
    ----------
    question : str
        The original natural-language question.
    sections : dict[str, str]
        Mapping ``section_title -> markdown_body``. Always includes at
        least: ``Question``, ``Data``, ``Identification``,
        ``Estimator``, ``Results``, ``Robustness``, ``References``.
    workflow : CausalWorkflow
        The underlying workflow object — exposes the raw fitted result
        (``workflow.result``), the diagnostics, the recommendation, etc.
    fmt : str
        Default output format (``markdown`` / ``tex`` / ``docx``).
    citations : list of str
        BibTeX-style entries collected from each estimator's ``cite()``.
    parsed_hints : dict
        What the question parser extracted, for transparency / debugging.
    degradations : list of dict
        Structured record of optional sub-steps that failed and were
        skipped (covariate balance, CI rendering, DAG appendix, citation
        extraction, provenance attachment, …).  Each entry has at least
        ``section``, ``error_type``, ``message``; some carry ``detail``.
        Empty when the draft is fully populated.  See
        :class:`statspai.workflow.WorkflowDegradedWarning`.
    """
    question: str
    sections: Dict[str, str]
    workflow: Any
    fmt: str
    citations: List[str] = field(default_factory=list)
    parsed_hints: Dict[str, Any] = field(default_factory=dict)
    dag: Any = None
    dag_treatment: Optional[str] = None
    dag_outcome: Optional[str] = None
    degradations: List[Dict[str, Any]] = field(default_factory=list)

    # ----- rendering -------------------------------------------------- #

    def to_markdown(self) -> str:
        order = [
            "Question", "Data", "Identification",
            "Estimator", "Results", "Robustness",
            "Reviewer Audit", "Pipeline notes", "Causal DAG", "References",
        ]
        chunks: List[str] = []
        for title in order:
            body = self.sections.get(title)
            if not body:
                continue
            chunks.append(f"## {title}\n\n{body.rstrip()}\n")
        # Append any extra sections in insertion order.
        for title, body in self.sections.items():
            if title not in order and body:
                chunks.append(f"## {title}\n\n{body.rstrip()}\n")
        return "\n".join(chunks)

    def to_tex(self) -> str:
        """Render to a LaTeX article skeleton.

        Each section becomes ``\\section{...}``; markdown bullet lists
        and code fences are translated to LaTeX equivalents.
        """
        body_lines: List[str] = []
        for title, body in self.sections.items():
            body_lines.append(f"\\section{{{_tex_escape(title)}}}")
            body_lines.append(_md_to_tex(body))
            body_lines.append("")
        bib = ""
        if self.citations:
            bib_items = "\n".join(f"\\bibitem{{r{i}}} {_tex_escape(c)}"
                                  for i, c in enumerate(self.citations))
            bib = (
                "\\begin{thebibliography}{99}\n"
                f"{bib_items}\n"
                "\\end{thebibliography}\n"
            )
        return (
            "\\documentclass{article}\n"
            "\\usepackage[utf8]{inputenc}\n"
            "\\usepackage{hyperref}\n"
            "\\title{Causal Analysis Draft}\n"
            "\\begin{document}\n"
            "\\maketitle\n\n"
            + "\n".join(body_lines)
            + "\n" + bib +
            "\\end{document}\n"
        )

    def to_docx(self, path: str) -> None:
        """Write a Word document to ``path``.

        Uses the workflow's already-fit result's ``to_docx`` if available;
        otherwise falls back to dropping a markdown file with a ``.docx``
        warning header (no python-docx hard dep).
        """
        try:
            import docx  # type: ignore
        except ImportError:
            # Fallback: write markdown to disk with a notice.
            with open(path, 'w', encoding='utf-8') as fh:
                fh.write(
                    "# (python-docx not installed; markdown fallback)\n\n"
                )
                fh.write(self.to_markdown())
            return
        doc = docx.Document()
        doc.add_heading("Causal Analysis Draft", level=0)
        for title, body in self.sections.items():
            doc.add_heading(title, level=1)
            for line in body.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
        doc.save(path)

    def to_qmd(
        self,
        *,
        title: str = "Causal Analysis Draft",
        author: Optional[str] = None,
        formats: Optional[List[str]] = None,
        bibliography: Optional[str] = None,
        csl: Optional[str] = None,
        include_provenance: bool = True,
    ) -> str:
        """Render to a Quarto (``.qmd``) document.

        Quarto is the multi-format manuscript default: a single source compiles
        to PDF / HTML / DOCX / Beamer with cross-refs, citations (CSL),
        and embedded code chunks. ``sp.paper()`` already produces all
        the prose; this method just wraps it in the correct YAML
        frontmatter so ``quarto render paper.qmd`` Just Works.

        Parameters
        ----------
        title : str
            ``title:`` field in the YAML frontmatter.
        author : str, optional
            ``author:``. When omitted, no author line is emitted (Quarto
            handles that fine).
        formats : list of str, optional
            Output formats Quarto should support. Default
            ``["pdf", "html", "docx"]`` covers the common journal
            workflows. Pass e.g. ``["pdf", "beamer"]`` for slide decks.
        bibliography : str, optional
            Path Quarto should resolve for citation lookup, e.g.
            ``"paper.bib"``. When omitted, the YAML omits the field
            entirely (so Quarto won't error if no .bib file exists);
            but if ``self.citations`` is non-empty we default to
            ``"paper.bib"`` because :func:`sp.replication_pack` writes
            citations there alongside the rendered draft.
        csl : str, optional
            CSL style file (e.g. ``"american-economic-association.csl"``).
            Pure pass-through.
        include_provenance : bool, default True
            Append a Reproducibility appendix with
            :func:`format_provenance` when ``self.workflow.result`` carries
            a ``_provenance`` record.

        Returns
        -------
        str
            The complete ``.qmd`` document as a single string.

        Notes
        -----
        - The body sections are the same as :meth:`to_markdown` —
          standard markdown with ``## H2`` headers, which Quarto will
          render natively.
        - Code chunks are *not* injected by default. When the calling
          script wants the ``.qmd`` to re-execute the analysis on each
          render, pass it through :func:`sp.replication_pack` which
          writes both the ``.qmd`` and a ``code/script.py`` reproducer.
        """
        formats = formats or ["pdf", "html", "docx"]
        bib_path = bibliography
        if bib_path is None and self.citations:
            bib_path = "paper.bib"

        yaml_lines: List[str] = ["---", f"title: {_yaml_str(title)}"]
        if author:
            yaml_lines.append(f"author: {_yaml_str(author)}")
        yaml_lines.append(f"date: \"{_dt.date.today().isoformat()}\"")
        if self.question:
            yaml_lines.append(
                f"subtitle: {_yaml_str(self.question)}"
            )
        # ``format:`` block.
        if len(formats) == 1:
            yaml_lines.append(f"format: {formats[0]}")
        else:
            yaml_lines.append("format:")
            for f in formats:
                yaml_lines.append(f"  {f}: default")
        if bib_path:
            yaml_lines.append(f"bibliography: {_yaml_str(bib_path)}")
        if csl:
            # Accept short journal names ('aer' / 'qje' / ...) and resolve
            # them to the canonical .csl filename. Pre-existing .csl
            # paths pass through untouched.  When resolution itself
            # blows up (corrupt csl_filename table, encoding error,
            # etc.) we still pass the user's raw input through, but
            # surface the resolver failure so they know the short-name
            # mapping was bypassed.
            try:
                from ..output._bibliography import csl_filename
                resolved = csl_filename(csl)
            except Exception as exc:
                record_degradation(
                    self.degradations,
                    section="CSL short-name resolution",
                    exc=exc,
                    detail=f"csl={csl!r}",
                )
                resolved = csl
            yaml_lines.append(f"csl: {_yaml_str(resolved)}")
        # Provenance into YAML for machine-readable traceability.
        prov = self._workflow_provenance()
        if include_provenance and prov is not None:
            yaml_lines.append("statspai:")
            yaml_lines.append(
                f"  version: \"{prov.statspai_version}\""
            )
            yaml_lines.append(f"  run_id: \"{prov.run_id}\"")
            if prov.data_hash:
                yaml_lines.append(
                    f"  data_hash: \"{prov.data_hash}\""
                )
        yaml_lines.append("---")
        yaml = "\n".join(yaml_lines)

        # Body — identical section ordering to to_markdown().
        order = [
            "Question", "Data", "Identification",
            "Estimator", "Results", "Robustness",
            "Reviewer Audit", "Pipeline notes", "Causal DAG", "References",
        ]
        # When self.dag is set, regenerate the Causal DAG body with the
        # Quarto-native mermaid block instead of the markdown text-art.
        sections_view = dict(self.sections)
        if self.dag is not None:
            sections_view["Causal DAG"] = _render_dag_section(
                self.dag,
                treatment=self.dag_treatment,
                outcome=self.dag_outcome,
                fmt="qmd",
                degradations=self.degradations,
            )
        chunks: List[str] = []
        for t in order:
            body = sections_view.get(t)
            if not body:
                continue
            chunks.append(f"## {t}\n\n{body.rstrip()}\n")
        for t, body in sections_view.items():
            if t not in order and body:
                chunks.append(f"## {t}\n\n{body.rstrip()}\n")

        # Reproducibility appendix.
        if include_provenance and prov is not None:
            chunks.append(
                "## Reproducibility {.appendix}\n\n"
                "```\n"
                f"{format_provenance(prov)}\n"
                "```\n"
            )

        return yaml + "\n\n" + "\n".join(chunks)

    def _workflow_provenance(self):
        wf = self.workflow
        if wf is None:
            return None
        result = getattr(wf, "result", None)
        if result is None:
            return None
        return get_provenance(result)

    def write(self, path: str) -> None:
        """Write the draft to disk in the format inferred from the path
        extension (``.md`` / ``.tex`` / ``.docx`` / ``.qmd``)."""
        lower = path.lower()
        if lower.endswith('.tex'):
            with open(path, 'w', encoding='utf-8') as fh:
                fh.write(self.to_tex())
        elif lower.endswith('.docx'):
            self.to_docx(path)
        elif lower.endswith('.qmd'):
            with open(path, 'w', encoding='utf-8') as fh:
                fh.write(self.to_qmd())
        else:
            with open(path, 'w', encoding='utf-8') as fh:
                fh.write(self.to_markdown())

    def summary(self) -> str:
        n = len(self.sections)
        return (
            "PaperDraft\n"
            "==========\n"
            f"  Question     : {self.question[:60]!r}"
            f"{'...' if len(self.question) > 60 else ''}\n"
            f"  Sections     : {n} ({', '.join(self.sections.keys())})\n"
            f"  Citations    : {len(self.citations)}\n"
            f"  Default fmt  : {self.fmt}\n"
        )

    def to_dict(self) -> Dict[str, Any]:
        return {
            'question': self.question,
            'sections': dict(self.sections),
            'parsed_hints': dict(self.parsed_hints),
            'citations': list(self.citations),
            'degradations': list(self.degradations),
            'fmt': self.fmt,
        }


# --------------------------------------------------------------------- #
#  Helpers
# --------------------------------------------------------------------- #


def _record_note(notes: Optional[List[str]], message: str) -> None:
    if notes is None:
        return
    note = str(message).strip()
    if note and note not in notes:
        notes.append(note)


def _notes_block(notes: List[str]) -> str:
    return "\n".join(f"- {note}" for note in notes)


def _render_dag_section(
    dag: Any,
    *,
    treatment: Optional[str] = None,
    outcome: Optional[str] = None,
    fmt: str = "markdown",
    degradations: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """Render a Causal DAG appendix section.

    Accepts a :class:`statspai.dag.graph.DAG` or any duck-typed object
    exposing ``nodes`` / ``edges`` / ``observed_nodes`` /
    ``adjustment_sets(t, y)`` / ``backdoor_paths(t, y)`` /
    ``bad_controls(t, y)``.

    For ``fmt='qmd'`` we emit a Quarto-native mermaid code block; for
    everything else (markdown / tex) we emit a plain text rendering
    that survives any downstream pipeline (LaTeX users typically
    typeset DAGs with TikZ; we do not attempt that here).

    Optional ``degradations`` collects per-bullet failures
    (adjustment_sets / backdoor_paths / bad_controls) — the outer
    section still renders, but the user gets a
    :class:`WorkflowDegradedWarning` and a structured trail when an
    individual back-door analysis can't be computed.
    """
    if dag is None:
        return ""

    nodes = sorted(getattr(dag, "observed_nodes", None) or
                   getattr(dag, "nodes", set()))
    edges = list(getattr(dag, "edges", []) or [])

    lines: List[str] = []

    # Mermaid block (Quarto renders natively; pandoc renders to SVG via
    # mermaid-cli when available; plain markdown viewers show the text).
    if fmt == "qmd":
        lines.append("```{mermaid}")
        lines.append("%%| fig-cap: Declared causal DAG")
        lines.append("graph LR")
        for u, v in edges:
            # Mermaid is sensitive to leading underscores in node ids;
            # latent ``_L_*`` nodes are quoted to avoid conflicts.
            uu = f'"{u}"' if u.startswith("_") else u
            vv = f'"{v}"' if v.startswith("_") else v
            lines.append(f"  {uu} --> {vv}")
        lines.append("```")
        lines.append("")

    # Variables list (always).
    if nodes:
        lines.append("**Variables**: " + ", ".join(f"`{n}`" for n in nodes))
        lines.append("")

    # Edges list (always — survives non-mermaid renderers).
    if edges:
        lines.append("**Edges**:")
        for u, v in edges:
            lines.append(f"- `{u}` → `{v}`")
        lines.append("")

    # Latent / unobserved confounders.
    latent = sorted(n for n in (getattr(dag, "nodes", set()) or set())
                    if isinstance(n, str) and n.startswith("_L_"))
    if latent:
        lines.append("**Latent common causes** (unobserved):")
        for n in latent:
            label = n.replace("_L_", "").replace("_", " ↔ ")
            lines.append(f"- `{n}` ({label})")
        lines.append("")

    # Identification analysis (back-door + bad controls), when treatment
    # and outcome are known.
    if treatment and outcome:
        try:
            adj = dag.adjustment_sets(treatment, outcome)
        except Exception as exc:
            record_degradation(
                degradations,
                section="DAG adjustment_sets sub-analysis",
                exc=exc,
                detail=f"treatment={treatment!r} outcome={outcome!r}",
            )
            adj = None
        if adj:
            lines.append(
                f"**Adjustment sets** (back-door criterion for "
                f"`{treatment}` → `{outcome}`):"
            )
            for s in adj:
                if not s:
                    lines.append("- ∅ (no controls needed)")
                else:
                    lines.append(
                        "- {" + ", ".join(f"`{x}`" for x in sorted(s)) + "}"
                    )
            lines.append("")
        try:
            bd = dag.backdoor_paths(treatment, outcome)
        except Exception as exc:
            record_degradation(
                degradations,
                section="DAG backdoor_paths sub-analysis",
                exc=exc,
                detail=f"treatment={treatment!r} outcome={outcome!r}",
            )
            bd = None
        if bd:
            lines.append(
                f"**Back-door paths** from `{treatment}` to `{outcome}`:"
            )
            for path in bd:
                arrow = " — ".join(f"`{x}`" for x in path)
                lines.append(f"- {arrow}")
            lines.append("")
        try:
            bad = dag.bad_controls(treatment, outcome)
        except Exception as exc:
            record_degradation(
                degradations,
                section="DAG bad_controls sub-analysis",
                exc=exc,
                detail=f"treatment={treatment!r} outcome={outcome!r}",
            )
            bad = None
        if bad:
            lines.append("**Bad controls** (do **not** condition on these):")
            for var, reason in bad.items():
                lines.append(f"- `{var}` — {reason}")
            lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def _yaml_str(value: str) -> str:
    """Quote a string safely for inclusion in a YAML scalar value.

    Always uses double-quotes and escapes any embedded ``"`` / ``\\``.
    Newlines are folded to a literal space (Quarto YAML headers don't
    play nicely with multi-line scalars in our context).
    """
    if value is None:
        return '""'
    s = str(value).replace("\n", " ").strip()
    # Backslash first, then double-quote.
    s = s.replace("\\", "\\\\").replace('"', '\\"')
    return f'"{s}"'


def _tex_escape(s: str) -> str:
    """Minimal LaTeX-escape for free-form text in section bodies."""
    if not isinstance(s, str):
        s = str(s)
    out = (s.replace("\\", r"\textbackslash{}")
           .replace("&", r"\&")
           .replace("%", r"\%")
           .replace("$", r"\$")
           .replace("#", r"\#")
           .replace("_", r"\_")
           .replace("{", r"\{")
           .replace("}", r"\}")
           .replace("~", r"\textasciitilde{}")
           .replace("^", r"\textasciicircum{}"))
    return out


def _md_to_tex(md: str) -> str:
    """Lightweight markdown → LaTeX translation for paper-section bodies.

    Handles: bold (**text**), bullet lists, fenced code blocks, inline
    code (`x`). Anything more elaborate falls through as escaped text.
    """
    out_lines: List[str] = []
    in_list = False
    in_code = False
    for ln in md.split("\n"):
        stripped = ln.rstrip()
        if stripped.startswith("```"):
            if in_code:
                out_lines.append(r"\end{verbatim}")
                in_code = False
            else:
                if in_list:
                    out_lines.append(r"\end{itemize}")
                    in_list = False
                out_lines.append(r"\begin{verbatim}")
                in_code = True
            continue
        if in_code:
            out_lines.append(stripped)
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            if not in_list:
                out_lines.append(r"\begin{itemize}")
                in_list = True
            item = stripped[2:].strip()
            out_lines.append(r"  \item " + _inline_md_to_tex(item))
            continue
        if in_list:
            out_lines.append(r"\end{itemize}")
            in_list = False
        if stripped == "":
            out_lines.append("")
        else:
            out_lines.append(_inline_md_to_tex(stripped))
    if in_list:
        out_lines.append(r"\end{itemize}")
    if in_code:
        out_lines.append(r"\end{verbatim}")
    return "\n".join(out_lines)


def _inline_md_to_tex(text: str) -> str:
    """Translate inline-markdown markers to LaTeX. Order matters: handle
    code spans before bold so we don't escape backticks inside code."""
    out = text
    # Code spans `...`
    out = re.sub(r"`([^`]+)`",
                 lambda m: r"\texttt{" + _tex_escape(m.group(1)) + "}",
                 out)
    # Bold **...**
    out = re.sub(r"\*\*([^*]+)\*\*",
                 lambda m: r"\textbf{" + _tex_escape(m.group(1)) + "}",
                 out)
    # Italic *...* (after bold so we don't double-match)
    out = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)",
                 lambda m: r"\emph{" + _tex_escape(m.group(1)) + "}",
                 out)
    # Anything left is plain text — escape special chars but leave
    # already-emitted LaTeX commands alone. The previous regexes have
    # produced LaTeX commands containing braces / backslashes; do a
    # cheap heuristic: only escape lines that don't already contain
    # ``\textbf`` / ``\texttt`` / ``\emph``.
    if (r"\textbf" not in out and r"\texttt" not in out
            and r"\emph" not in out):
        out = _tex_escape(out)
    return out


def _eda_block(data: pd.DataFrame, y: Optional[str],
               treatment: Optional[str],
               covariates: Optional[List[str]],
               degradations: Optional[List[Dict[str, Any]]] = None) -> str:
    """Build a brief EDA markdown section (size, balance, missingness).

    Optional ``degradations`` list, when supplied, receives a structured
    record (and a :class:`WorkflowDegradedWarning` fires) whenever a
    sub-section has to be skipped — e.g. the covariate-balance table on
    a non-numeric covariate.  Pass the owning ``PaperDraft.degradations``
    list so the caller can introspect what got dropped.
    """
    lines: List[str] = []
    n_rows, n_cols = data.shape
    lines.append(f"- Sample size: **{n_rows:,}** rows, **{n_cols}** columns.")
    miss = data.isna().mean()
    miss = miss[miss > 0].sort_values(ascending=False)
    if not miss.empty:
        lines.append("- Missingness (top 5):")
        for col, frac in miss.head(5).items():
            lines.append(f"    - `{col}`: {frac*100:.1f}%")
    else:
        lines.append("- Missingness: none detected in the analysis frame.")
    if y and y in data.columns:
        ys = data[y].dropna()
        if pd.api.types.is_numeric_dtype(ys):
            lines.append(
                f"- Outcome `{y}`: "
                f"mean={ys.mean():.3f}, sd={ys.std():.3f}, "
                f"median={ys.median():.3f}, n={len(ys)}."
            )
    if treatment and treatment in data.columns:
        tr = data[treatment].dropna()
        if tr.nunique() <= 10:
            counts = tr.value_counts().sort_index()
            shares = (counts / counts.sum() * 100).round(1)
            lines.append(
                f"- Treatment `{treatment}` distribution: "
                + ", ".join(
                    f"{int(k) if k == int(k) else k}="
                    f"{int(counts.loc[k])} ({shares.loc[k]}%)"
                    for k in counts.index
                )
            )
        else:
            lines.append(
                f"- Treatment `{treatment}` (continuous): "
                f"mean={tr.mean():.3f}, sd={tr.std():.3f}."
            )
    # Optional covariate balance — only for binary treatment.
    if (treatment and treatment in data.columns
            and covariates and len(covariates) <= 8
            and data[treatment].nunique() == 2):
        balance_lines: List[str] = []
        try:
            grp = data.groupby(treatment)[covariates].mean()
            if grp.shape[0] == 2:
                balance_lines.append("")
                balance_lines.append("Mean covariates by treatment arm:")
                balance_lines.append("")
                balance_lines.append("| covariate | "
                                     + " | ".join(str(g) for g in grp.index)
                                     + " | std-diff |")
                balance_lines.append(
                    "|---|" + "|".join(["---"] * grp.shape[0]) + "|---|"
                )
                pooled_std = data[covariates].std()
                vals0 = grp.iloc[0]
                vals1 = grp.iloc[1]
                std_diff = (vals1 - vals0) / pooled_std.replace(0, np.nan)
                for c in covariates:
                    balance_lines.append(
                        f"| {c} | {grp.iloc[0][c]:.3f} "
                        f"| {grp.iloc[1][c]:.3f} "
                        f"| {std_diff[c]:.3f} |"
                    )
        except Exception as exc:
            # Most common cause: non-numeric covariate slipped into the
            # list, or all-NaN column. Fail loud (per CLAUDE.md §3.7),
            # drop the partial table, and let the rest of the EDA stand.
            record_degradation(
                degradations,
                section="EDA covariate balance table",
                exc=exc,
                detail=f"covariates={list(covariates)}",
            )
        else:
            lines.extend(balance_lines)
    return "\n".join(lines)


def _section_from_workflow(
    workflow,
    *,
    include_robustness: bool = True,
    degradations: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, str]:
    """Extract Identification / Estimator / Results / Robustness sections
    from a fitted CausalWorkflow.

    Optional ``degradations`` collects structured records of any
    sub-section that had to fall back (e.g. CI rendering on a non-2-tuple
    ``ci`` attribute, or serialisation of a coefficient table that
    surfaces in an unexpected shape).  Pairs with
    :class:`WorkflowDegradedWarning` so callers see the failure surface.
    """
    sections: Dict[str, str] = {}

    # Identification
    diag = workflow.diagnostics
    lines: List[str] = []
    if diag is None:
        lines.append("No identification report available.")
    else:
        try:
            lines.append(f"**Verdict**: {diag.verdict}")
            lines.append("")
            if diag.findings:
                for f in diag.findings:
                    lines.append(f"- [{f.severity.upper()}] *{f.category}* — "
                                 f"{f.message}")
                    if f.suggestion:
                        lines.append(f"    - Fix: {f.suggestion}")
            else:
                lines.append("No identification issues flagged.")
        except Exception as exc:
            lines = [
                "Identification report available but not fully serialisable; "
                "see `paper.workflow.diagnostics`."
            ]
            record_degradation(
                degradations,
                section="Identification section serialisation",
                exc=exc,
                detail=f"diagnostics_type={type(diag).__name__}",
            )
    sections["Identification"] = "\n".join(lines)

    # Estimator
    rec = workflow.recommendation
    lines = []
    if rec is not None and getattr(rec, "recommendations", None):
        try:
            top = rec.recommendations[0]
            lines.append(f"- **Method**: {top['method']}")
            lines.append(f"- **Function**: `sp.{top['function']}()`")
            if top.get('reason'):
                lines.append(f"- **Rationale**: {top['reason']}")
            if top.get('assumptions'):
                lines.append("- **Key assumptions**: "
                             + ", ".join(top['assumptions']))
        except Exception as exc:
            lines.append(
                "Estimator recommendation available but not fully "
                "serialisable; see `paper.workflow.recommendation`."
            )
            record_degradation(
                degradations,
                section="Estimator section serialisation",
                exc=exc,
                detail=f"recommendation_type={type(rec).__name__}",
            )
    else:
        lines.append("No estimator recommendation produced.")
    sections["Estimator"] = "\n".join(lines)

    # Results
    r = workflow.result
    lines = []
    if r is not None and hasattr(r, 'estimate') and hasattr(r, 'se'):
        try:
            est = float(r.estimate)
            se = float(r.se)
            lines.append(
                f"- **{getattr(r, 'estimand', 'Effect')}**: "
                f"{est:.4f} (SE = {se:.4f})"
            )
            ci = getattr(r, 'ci', None)
            if ci is not None and not isinstance(ci, (pd.DataFrame, pd.Series)):
                try:
                    lo, hi = float(ci[0]), float(ci[1])
                    lines.append(f"- **95% CI**: [{lo:.4f}, {hi:.4f}]")
                except Exception as exc:
                    record_degradation(
                        degradations,
                        section="Results 95% CI rendering",
                        exc=exc,
                        detail=f"ci={ci!r}",
                    )
            pv = getattr(r, 'pvalue', np.nan)
            if pd.notna(pv):
                lines.append(f"- **p-value**: {float(pv):.4f}")
            n_obs = getattr(r, 'n_obs', None)
            if n_obs is not None:
                lines.append(f"- **N obs**: {int(n_obs)}")
        except Exception as exc:
            lines.append("Result available but not fully serialisable; "
                         "see `paper.workflow.result`.")
            record_degradation(
                degradations,
                section="Results section serialisation (estimate/se path)",
                exc=exc,
                detail=f"result_type={type(r).__name__}",
            )
    elif r is not None and hasattr(r, 'params'):
        try:
            main = (workflow.treatment or list(r.params.index)[0])
            if main in r.params.index:
                est = float(r.params[main])
                se = float(r.std_errors[main])
                lines.append(f"- **{main}**: {est:.4f} (SE = {se:.4f})")
            else:
                lines.append("Coefficient table available; see "
                             "`paper.workflow.result.params`.")
        except Exception as exc:
            lines.append("Result available but not fully serialisable.")
            record_degradation(
                degradations,
                section="Results section serialisation (params path)",
                exc=exc,
                detail=f"result_type={type(r).__name__}",
            )
    else:
        lines.append("No fitted result available.")
    sections["Results"] = "\n".join(lines)

    if include_robustness:
        # Robustness — prefer the structured ``RobustnessReport`` rendered
        # by the shared battery (severity icons, plain-English interpretation,
        # design-specific ordering).  Fall back to the legacy flat-key
        # rendering only when an older caller assigned ``robustness_findings``
        # directly without going through ``CausalWorkflow.robustness()``.
        report = getattr(workflow, "_robustness_report", None)
        rendered = None
        if report is not None and not report.is_empty():
            try:
                rendered = report.to_markdown()
            except Exception as exc:
                record_degradation(
                    degradations,
                    section="Robustness section rendering",
                    exc=exc,
                    detail=f"report_type={type(report).__name__}",
                )
        if rendered is None:
            findings = workflow.robustness_findings or {}
            # Strip the ``_findings`` / ``_design`` / ``_notes`` private
            # keys that the new ``to_dict()`` shape adds — they belong in
            # the structured payload, not the human-readable bullet list.
            flat = {
                k: v for k, v in findings.items()
                if not str(k).startswith("_")
            }
            lines = []
            if flat:
                for k, v in flat.items():
                    if isinstance(v, (int, float, np.integer, np.floating)):
                        if isinstance(v, (int, np.integer)):
                            lines.append(f"- {k.replace('_', ' ').title()}: "
                                         f"{int(v)}")
                        else:
                            lines.append(f"- {k.replace('_', ' ').title()}: "
                                         f"{float(v):.4f}")
                    elif isinstance(v, dict):
                        lines.append(f"- {k.replace('_', ' ').title()}:")
                        for kk, vv in list(v.items())[:8]:
                            lines.append(f"    - {kk}: {vv}")
                    else:
                        lines.append(f"- {k.replace('_', ' ').title()}: {v}")
            else:
                lines.append("No robustness findings produced.")
            rendered = "\n".join(lines)
        sections["Robustness"] = rendered

    return sections


def _reviewer_audit_section(
    *,
    workflow: Any = None,
    result: Any = None,
    estimator: Optional[str] = None,
    degradations: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """Build a reviewer-facing audit section for a fitted draft."""
    lines: List[str] = []
    lines.append("**Reviewer-mode audit**")
    lines.append("")

    # Estimator registry evidence.
    est_name = estimator
    if not est_name and workflow is not None:
        rec = getattr(workflow, "recommendation", None)
        try:
            if rec is not None and rec.recommendations:
                est_name = rec.recommendations[0].get("function")
        except Exception:
            est_name = None
    if est_name:
        try:
            from ..registry import describe_function
            spec = describe_function(est_name)
            lines.append(
                f"- **Registry**: `sp.{est_name}` is "
                f"`stability={spec.get('stability')}`, "
                f"`validation_status={spec.get('validation_status')}`."
            )
            notes = spec.get("validation_notes") or []
            if notes:
                shown = "; ".join(str(n) for n in notes[:3])
                lines.append(f"- **Validation evidence**: {shown}.")
            limitations = spec.get("limitations") or []
            if limitations:
                lines.append("- **Known implementation limitations**:")
                lines.extend(f"    - {lim}" for lim in limitations[:5])
        except Exception as exc:
            lines.append(
                f"- **Registry**: lookup for `sp.{est_name}` failed "
                f"({type(exc).__name__}: {exc})."
            )

    # Identification report.
    diag = getattr(workflow, "diagnostics", None) if workflow is not None else None
    if diag is not None:
        verdict = getattr(diag, "verdict", None)
        if verdict:
            lines.append(f"- **Identification verdict**: `{verdict}`.")
        findings = getattr(diag, "findings", None) or []
        if findings:
            lines.append("- **Identification findings to defend**:")
            for finding in findings[:6]:
                sev = getattr(finding, "severity", "info")
                msg = getattr(finding, "message", str(finding))
                lines.append(f"    - [{str(sev).upper()}] {msg}")

    # Post-estimation contract.
    target = result or getattr(workflow, "result", None)
    if target is not None:
        try:
            from ..postestimation import postestimation_contract
            contract = postestimation_contract(target)
            available = ", ".join(sorted(contract["available"].keys())[:10])
            lines.append(f"- **Post-estimation surface**: {available}.")
        except Exception as exc:
            lines.append(
                f"- **Post-estimation surface**: unavailable "
                f"({type(exc).__name__}: {exc})."
            )

        violations = getattr(target, "violations", None)
        if callable(violations):
            try:
                v = violations()
                if v:
                    lines.append("- **Result violations**:")
                    if isinstance(v, dict):
                        iterable = list(v.items())[:8]
                        lines.extend(f"    - `{k}`: {val}" for k, val in iterable)
                    else:
                        lines.append(f"    - {v}")
                else:
                    lines.append("- **Result violations**: none reported by result.")
            except Exception as exc:
                lines.append(
                    f"- **Result violations**: check failed "
                    f"({type(exc).__name__}: {exc})."
                )

        prov = get_provenance(target)
        if prov is not None:
            lines.append(
                f"- **Provenance**: run `{prov.run_id}`, "
                f"data hash `{prov.data_hash or 'not recorded'}`."
            )
        else:
            lines.append("- **Provenance**: no attached provenance record found.")

    if degradations:
        lines.append("- **Pipeline degradations**:")
        for item in degradations[:8]:
            detail = f" ({item.get('detail')})" if item.get("detail") else ""
            lines.append(
                f"    - {item.get('section')}: {item.get('error_type')}: "
                f"{item.get('message')}{detail}"
            )

    lines.append("")
    lines.append("**Reviewer checklist**")
    lines.append("- Re-run the replication script or `sp.replication_pack()` output.")
    lines.append("- Check identification assumptions against the study design, not only the code.")
    lines.append("- Inspect overlap, pre-trends, weak instruments, or bandwidth sensitivity when relevant.")
    lines.append("- Confirm exported tables are generated from `tidy()`/`glance()`/`sp.collect()` artifacts.")
    return "\n".join(lines)


# --------------------------------------------------------------------- #
#  Top-level entry point
# --------------------------------------------------------------------- #


def paper(
    data,
    question: Optional[str] = None,
    *,
    y: Optional[str] = None,
    treatment: Optional[str] = None,
    covariates: Optional[List[str]] = None,
    id: Optional[str] = None,
    time: Optional[str] = None,
    running_var: Optional[str] = None,
    instrument: Optional[str] = None,
    cutoff: Optional[float] = None,
    cohort: Optional[str] = None,
    cluster: Optional[str] = None,
    design: Optional[str] = None,
    dag=None,
    llm: Optional[str] = None,
    llm_client: Any = None,
    llm_domain: str = "",
    fmt: str = 'markdown',
    output_path: Optional[str] = None,
    include_eda: bool = True,
    include_robustness: bool = True,
    cite: bool = True,
    strict: bool = False,
    # v1.13: forwarded to sp.causal -> sp.recommend; default False
    # keeps frontier MVP estimators out of auto-generated drafts.
    allow_experimental: bool = False,
    reviewer_mode: bool = False,
) -> PaperDraft:
    """End-to-end "data → publication-draft" pipeline.

    Run :class:`CausalWorkflow` and assemble its outputs into a
    structured :class:`PaperDraft` (markdown / LaTeX / Word).

    Parameters
    ----------
    data : pd.DataFrame
        Analysis frame.
    question : str
        Natural-language causal question, e.g.
        ``"effect of training on wages"``. Used both to seed the
        question section of the draft and to fill in any missing
        column hints (``treatment`` / ``y``) when not given explicitly.
    y, treatment, covariates, id, time, running_var, instrument, cutoff, cohort, cluster, design, dag : optional
        Forwarded to :func:`sp.causal`. When omitted, the question
        parser tries to infer them from ``question``.
    fmt : {'markdown', 'tex', 'docx'}, default 'markdown'
        Default rendering format. The :class:`PaperDraft` always knows
        how to emit each format on demand via its ``.to_*()`` methods.
    output_path : str, optional
        When provided, write the rendered draft to disk in the format
        inferred from the path extension (``.md`` / ``.tex`` /
        ``.docx``).
    include_eda : bool, default True
        Include the Data section (descriptives + balance).
    include_robustness : bool, default True
        Include the Robustness section.
    cite : bool, default True
        Pull bibliography entries from the fitted result's ``cite()``
        method (when available).
    strict : bool, default False
        Forwarded to :func:`sp.causal` — when True, identification
        warnings escalate to errors.
    allow_experimental : bool, default False
        Forwarded to :func:`sp.causal` → :func:`sp.recommend`. When
        ``False`` (the agent-safe default), auto-generated drafts
        cannot land on a function whose registry entry is
        ``stability='experimental'`` or ``'deprecated'``. Set ``True``
        when you are explicitly drafting a paper around frontier
        methods (e.g. ``causal_text`` or ``did_multiplegt_dyn``); the
        Pipeline notes section records what was filtered.
    reviewer_mode : bool, default False
        Add a "Reviewer Audit" section summarizing registry validation
        status, identification findings, post-estimation capabilities,
        provenance, violations, and a replication checklist.

    Returns
    -------
    PaperDraft

    Examples
    --------
    >>> import statspai as sp
    >>> draft = sp.paper(df, "effect of training on wages", design='did',
    ...                  treatment='trained', y='wage', time='year',
    ...                  id='worker_id')
    >>> print(draft.to_markdown()[:500])
    >>> draft.write('analysis.tex')

    Notes
    -----
    The question parser is purely additive — explicit kwargs always win.
    Pass everything you know; the parser fills in only what's missing.
    """
    if fmt not in {'markdown', 'tex', 'docx', 'qmd'}:
        raise ValueError(
            f"Unknown fmt={fmt!r}. Use 'markdown', 'tex', 'docx', or 'qmd'."
        )

    # Estimand-first dispatch: if the first positional arg is a
    # CausalQuestion (already declared population/estimand/design/data),
    # route through identify() + estimate() rather than the
    # natural-language workflow path. This is the v1.7 estimand-first
    # entry point — the Target-Trial-Protocol-shaped declaration drives
    # the analysis end-to-end.
    try:
        from ..question.question import CausalQuestion as _CQ
    except Exception:
        _CQ = None  # pragma: no cover — package-internal
    if _CQ is not None and isinstance(data, _CQ):
        if question is None:
            question = (
                f"effect of {data.treatment} on {data.outcome}"
                + (f" in {data.population}" if data.population else "")
            )
        return paper_from_question(
            data, question=question, fmt=fmt,
            output_path=output_path,
            include_robustness=include_robustness,
            cite=cite, dag=dag,
            reviewer_mode=reviewer_mode,
        )

    cols = list(data.columns)
    parsed = parse_question(question or "", cols)

    # Explicit args win; parser fills gaps.
    y_eff = y or parsed.get('y')
    t_eff = treatment or parsed.get('treatment')
    design_eff = design or parsed.get('design')
    instrument_eff = instrument or parsed.get('instrument')
    running_var_eff = running_var or parsed.get('running_var')
    cutoff_eff = cutoff if cutoff is not None else parsed.get('cutoff')

    if y_eff is None:
        raise ValueError(
            "Could not determine the outcome `y`. Pass `y=...` explicitly "
            "or include 'effect of X on Y' in the question."
        )

    # Track optional sub-steps that fall back (CLAUDE.md §3.7).  Declared
    # early so the LLM-DAG and provenance blocks below can record into
    # the same list.
    degradations: List[Dict[str, Any]] = []

    # LLM-DAG auto-propose: when ``llm`` is requested and the user
    # didn't pass an explicit ``dag``, ask the LLM (or fall back to the
    # deterministic heuristic) to propose one. Resolution of provider
    # / model / API key goes through the layered ``get_llm_client``
    # — see ``statspai.causal_llm._resolver``.
    if dag is None and llm:
        try:
            from ..causal_llm._resolver import get_llm_client
            from ..causal_llm.llm_dag import llm_dag_propose
            from ..dag.graph import DAG as _DAG
            client = llm_client
            # llm='auto' means "try to resolve a real client; fall back
            # to the heuristic if none can be resolved". llm='heuristic'
            # forces the offline backend regardless.
            if client is None and str(llm).lower() != "heuristic":
                try:
                    client = get_llm_client(allow_interactive=False)
                except Exception as resolver_exc:
                    # Hard error from resolver → fall back to heuristic
                    # rather than blowing up the whole paper pipeline.
                    record_degradation(
                        degradations,
                        section="LLM client resolution (auto-DAG)",
                        exc=resolver_exc,
                        detail=f"llm={llm!r}",
                    )
                    client = None
            proposal = llm_dag_propose(
                variables=cols,
                domain=llm_domain or (question or ""),
                client=client,
            )
            if proposal.edges:
                spec = "; ".join(f"{u} -> {v}" for u, v in proposal.edges)
                dag = _DAG(spec)
        except Exception as exc:
            # Failed to propose a DAG at all → degrade to no-DAG paper
            # but tell the user why (rather than silently producing a
            # paper without an Identification DAG they had asked for).
            record_degradation(
                degradations,
                section="auto-DAG proposal (sp.paper(llm=...))",
                exc=exc,
                detail=f"llm={llm!r}",
            )
            dag = None

    from .causal_workflow import causal as _causal
    workflow = _causal(
        data,
        y=y_eff,
        treatment=t_eff,
        covariates=covariates,
        id=id,
        time=time,
        running_var=running_var_eff,
        instrument=instrument_eff,
        cutoff=cutoff_eff,
        cohort=cohort,
        cluster=cluster,
        design=design_eff,
        dag=dag,
        strict=strict,
        auto_run=False,
        allow_experimental=allow_experimental,
    )
    # Drive the pipeline through to robustness, swallowing per-stage
    # failures into per-section fallback notes (the draft must always
    # produce something — agents shouldn't see a hard crash for one
    # bad estimator choice).
    pipeline_errors: List[str] = []
    for stage in ("diagnose", "recommend", "estimate"):
        try:
            getattr(workflow, stage)()
        except Exception as exc:  # pragma: no cover (defensive)
            pipeline_errors.append(
                f"`{stage}()` failed: {type(exc).__name__}: {exc}"
            )
            break
    if include_robustness:
        if workflow.result is not None:
            try:
                workflow.robustness()
            except Exception as exc:  # pragma: no cover (defensive)
                pipeline_errors.append(
                    f"`robustness()` failed: {type(exc).__name__}: {exc}"
                )
        else:
            pipeline_errors.append(
                "`robustness()` skipped because no fitted result was available."
            )

    # Attach provenance to the workflow's result so downstream
    # ``replication_pack`` / Quarto appendix / table footers can pick
    # it up. Estimators that wire their own ``attach_provenance`` at
    # ``fit()`` end already populate ``_provenance``; ``overwrite=False``
    # preserves their (more specific) record.
    try:
        from ..output._lineage import attach_provenance as _attach_prov
        if workflow.result is not None:
            _attach_prov(
                workflow.result,
                function=f"sp.causal[{workflow.design or 'auto'}]",
                params={
                    "y": y_eff,
                    "treatment": t_eff,
                    "design": design_eff or workflow.design,
                    "instrument": instrument_eff,
                    "running_var": running_var_eff,
                    "cutoff": cutoff_eff,
                    "covariates": covariates,
                },
                data=data,
                overwrite=False,
            )
    except Exception as exc:  # provenance loss is loud, not silent
        record_degradation(
            degradations,
            section="provenance attachment (workflow path)",
            exc=exc,
            detail=f"design={workflow.design or 'auto'}",
        )

    sections: Dict[str, str] = {}

    # Question
    sections["Question"] = (
        f"> {question.strip() if question else '(no question supplied)'}\n\n"
        f"- **Outcome**: `{y_eff}`\n"
        + (f"- **Treatment**: `{t_eff}`\n" if t_eff else "")
        + (f"- **Design (auto-detected)**: `{workflow.design}`\n"
           if workflow.design else "")
    )

    # Data / EDA
    if include_eda:
        sections["Data"] = _eda_block(
            data, y_eff, t_eff, covariates, degradations=degradations,
        )

    # Identification / Estimator / Results / Robustness
    sections.update(
        _section_from_workflow(
            workflow,
            include_robustness=include_robustness,
            degradations=degradations,
        )
    )

    # Causal DAG appendix (when the user passes a DAG).
    if dag is not None:
        try:
            sections["Causal DAG"] = _render_dag_section(
                dag, treatment=t_eff, outcome=y_eff, fmt="markdown",
                degradations=degradations,
            )
        except Exception as exc:
            record_degradation(
                degradations,
                section="causal DAG appendix",
                exc=exc,
                detail=f"dag_type={type(dag).__name__}",
            )

    # References — citation extraction is a §10 ("引用零幻觉") concern in
    # reverse: a silently dropped real citation is just as harmful as a
    # hallucinated one.  Surface the failure rather than swallow it.
    citations: List[str] = []
    if cite and workflow.result is not None:
        cite_fn = getattr(workflow.result, 'cite', None)
        if callable(cite_fn):
            try:
                ref = cite_fn()
                if ref:
                    citations.append(str(ref))
            except Exception as exc:
                record_degradation(
                    degradations,
                    section="citation extraction (workflow.result.cite)",
                    exc=exc,
                    detail=f"result_type={type(workflow.result).__name__}",
                )
    references_body = (
        "\n".join(f"- {c}" for c in citations)
        if citations else "_(No explicit citations attached — see "
        "`workflow.result.cite()` if available.)_"
    )

    if reviewer_mode:
        sections["Reviewer Audit"] = _reviewer_audit_section(
            workflow=workflow,
            result=workflow.result,
            degradations=degradations,
        )
    sections["References"] = references_body

    pipeline_notes: List[str] = []
    for error in pipeline_errors:
        _record_note(pipeline_notes, error)
    for note in getattr(workflow, "pipeline_notes", []) or []:
        _record_note(pipeline_notes, note)
    for item in degradations:
        detail = f" ({item['detail']})" if item.get("detail") else ""
        _record_note(
            pipeline_notes,
            f"{item['section']} degraded: {item['error_type']}: "
            f"{item['message']}{detail}",
        )
    if pipeline_notes:
        sections["Pipeline notes"] = _notes_block(pipeline_notes)

    draft = PaperDraft(
        question=question or "",
        sections=sections,
        workflow=workflow,
        fmt=fmt,
        citations=citations,
        parsed_hints=parsed,
        dag=dag,
        dag_treatment=t_eff,
        dag_outcome=y_eff,
        degradations=degradations,
    )

    if output_path is not None:
        draft.write(output_path)

    return draft


# --------------------------------------------------------------------- #
#  Estimand-first paper builder (CausalQuestion → PaperDraft)
# --------------------------------------------------------------------- #


def paper_from_question(
    q,
    *,
    question: Optional[str] = None,
    fmt: str = "markdown",
    output_path: Optional[str] = None,
    include_robustness: bool = True,
    cite: bool = True,
    dag: Any = None,
    reviewer_mode: bool = False,
) -> PaperDraft:
    """Build a :class:`PaperDraft` from a :class:`CausalQuestion`.

    The estimand-first entry point: instead of inferring the design
    from natural language and a DataFrame, the user *declares* the
    Target Trial Protocol (treatment / outcome / population / estimand
    / design / time-structure) up front, and this builder routes
    through ``q.identify()`` + ``q.estimate()`` to assemble a paper
    whose Methods / Results sections match what was pre-registered.

    Sections produced
    -----------------
    - **Question** — population, treatment, outcome, estimand,
      time-structure (verbatim from the declaration).
    - **Identification** — the IdentificationPlan's story + assumptions.
    - **Estimator** — chosen primary estimator + fallbacks.
    - **Results** — point estimate, SE, 95% CI, n.
    - **Robustness** — placeholder unless the user attaches a
      :class:`CausalWorkflow` (or rolls their own follow-up).
    - **References** — pulled from ``result.underlying.cite()`` when
      available.

    Parameters
    ----------
    q : CausalQuestion
        Declared causal question. Must have ``data`` set.
    question : str, optional
        Free-form text to embed in the Question section. Defaults to
        ``"effect of <treatment> on <outcome>"``.
    fmt : {'markdown', 'tex', 'docx', 'qmd'}, default 'markdown'
        Default rendering format on the returned ``PaperDraft``.
    output_path : str, optional
        Write the rendered draft to disk in the format inferred from
        the path extension.
    include_robustness : bool, default True
        Reserve the Robustness section (will say "not run" until the
        user attaches a workflow).
    cite : bool, default True
        Pull bibliography entries from ``result.underlying.cite()``
        when available.
    dag : object, optional
        Pre-built ``sp.dag`` graph. When provided, the draft's
        Identification section gains a *Causal DAG* subsection (text
        rendering for markdown / mermaid for qmd).
    reviewer_mode : bool, default False
        Add a reviewer-facing audit section with registry validation,
        post-estimation capabilities, provenance, and replication checks.

    Returns
    -------
    PaperDraft

    Examples
    --------
    >>> import statspai as sp
    >>> q = sp.causal_question(
    ...     "trained", "wage", data=df, design="did",
    ...     time="year", id="worker_id",
    ...     covariates=["edu"]
    ... )
    >>> draft = sp.paper(q, fmt='qmd')
    >>> draft.write("paper.qmd")
    """
    if q.data is None:
        raise ValueError(
            "CausalQuestion.data must be set before paper_from_question(). "
            "Pass data= when constructing the question."
        )
    if fmt not in {"markdown", "tex", "docx", "qmd"}:
        raise ValueError(
            f"Unknown fmt={fmt!r}. Use 'markdown', 'tex', 'docx', or 'qmd'."
        )

    if q._plan is None:
        plan = q.identify()
    else:
        plan = q._plan
    if q._result is None:
        result = q.estimate()
    else:
        result = q._result

    # Track optional sub-steps that fall back (CLAUDE.md §3.7).
    degradations: List[Dict[str, Any]] = []

    sections: Dict[str, str] = {}
    eff_question = question or (
        f"effect of {q.treatment} on {q.outcome}"
    )

    # Question section
    bits = [
        f"> {eff_question}",
        "",
        f"- **Population**: {q.population or '_(not specified)_'}",
        f"- **Treatment**: `{q.treatment}`",
        f"- **Outcome**: `{q.outcome}`",
        f"- **Estimand**: {plan.estimand}",
        f"- **Design**: `{q.design}`",
        f"- **Time structure**: {q.time_structure}",
    ]
    if q.covariates:
        bits.append(
            "- **Covariates**: " + ", ".join(f"`{c}`" for c in q.covariates)
        )
    if q.instruments:
        bits.append(
            "- **Instruments**: " + ", ".join(f"`{c}`" for c in q.instruments)
        )
    if q.cutoff is not None:
        bits.append(f"- **Cutoff**: {q.cutoff}")
    if q.notes:
        bits.append(f"- **Notes**: {q.notes}")
    sections["Question"] = "\n".join(bits)

    # Data section — declared frame stats.
    data = q.data
    n_rows, n_cols = data.shape
    miss = data.isna().mean()
    miss = miss[miss > 0].sort_values(ascending=False)
    data_lines = [
        f"- Sample size: **{n_rows:,}** rows, **{n_cols}** columns.",
    ]
    if not miss.empty:
        data_lines.append("- Missingness (top 5):")
        for col, frac in miss.head(5).items():
            data_lines.append(f"    - `{col}`: {frac*100:.1f}%")
    else:
        data_lines.append("- Missingness: none detected in the analysis frame.")
    if q.outcome in data.columns:
        ys = data[q.outcome].dropna()
        if pd.api.types.is_numeric_dtype(ys):
            data_lines.append(
                f"- Outcome `{q.outcome}`: mean={ys.mean():.3f}, "
                f"sd={ys.std():.3f}, median={ys.median():.3f}, n={len(ys)}."
            )
    sections["Data"] = "\n".join(data_lines)

    # Identification section
    id_lines = [plan.identification_story, ""]
    id_lines.append("**Required assumptions** (must defend):")
    for a in plan.assumptions:
        id_lines.append(f"- {a}")
    if plan.warnings:
        id_lines.append("")
        id_lines.append("**Warnings:**")
        for w in plan.warnings:
            id_lines.append(f"- {w}")
    if dag is not None:
        id_lines.append("")
        id_lines.append("**Causal DAG**: see appendix.")
    sections["Identification"] = "\n".join(id_lines)

    # Estimator section
    est_lines = [
        f"- **Primary estimator**: `sp.{plan.estimator}`",
        f"- **Estimand**: {plan.estimand}",
    ]
    if plan.fallback_estimators:
        est_lines.append(
            "- **Fallbacks** (if primary's assumptions fail): "
            + ", ".join(f"`sp.{f}`" for f in plan.fallback_estimators)
        )
    sections["Estimator"] = "\n".join(est_lines)

    # Results section
    lo, hi = result.ci
    res_lines = [
        f"- **{plan.estimand}** (via `sp.{plan.estimator}`): "
        f"**{result.estimate:+.4f}** (SE = {result.se:.4f})",
        f"- **95% CI**: [{lo:+.4f}, {hi:+.4f}]",
        f"- **N obs**: {int(result.n)}",
    ]
    sections["Results"] = "\n".join(res_lines)

    # Robustness — delegate to the shared battery so the
    # estimand-first path delivers the same diagnostic content as the
    # natural-language path.  ``run_robustness_battery`` never raises;
    # any per-check failure becomes a ``severity='check_failed'``
    # finding rather than aborting the section.
    if include_robustness:
        from ._robustness import run_robustness_battery
        # ``CausalQuestion.estimate`` returns an ``EstimationResult``
        # wrapping the underlying estimator result; the battery wants
        # the underlying object so it can introspect ``.violations()``,
        # ``.model_info``, etc.
        underlying = getattr(result, "underlying", None) or result
        try:
            report = run_robustness_battery(
                underlying,
                design=q.design,
                data=q.data,
                treatment=q.treatment,
                outcome=q.outcome,
                covariates=list(q.covariates) if q.covariates else None,
            )
            sections["Robustness"] = report.to_markdown()
        except Exception as exc:
            record_degradation(
                degradations,
                section="robustness battery (estimand-first path)",
                exc=exc,
                detail=f"design={q.design!r}",
            )

    # References — see paper()/citation note: a silently dropped citation
    # is just as harmful as a hallucinated one (§10).
    citations: List[str] = []
    if cite and result.underlying is not None:
        cite_fn = getattr(result.underlying, "cite", None)
        if callable(cite_fn):
            try:
                ref = cite_fn()
                if ref:
                    citations.append(str(ref))
            except Exception as exc:
                record_degradation(
                    degradations,
                    section="citation extraction (result.underlying.cite)",
                    exc=exc,
                    detail=f"underlying_type={type(result.underlying).__name__}",
                )
    references_body = (
        "\n".join(f"- {c}" for c in citations)
        if citations
        else "_(No explicit citations attached — see "
             "`result.underlying.cite()` if available.)_"
    )

    if reviewer_mode:
        target = result.underlying if result.underlying is not None else result
        sections["Reviewer Audit"] = _reviewer_audit_section(
            workflow=None,
            result=target,
            estimator=plan.estimator,
            degradations=degradations,
        )
    sections["References"] = references_body

    # DAG appendix when the user attached one.
    if dag is not None:
        try:
            sections["Causal DAG"] = _render_dag_section(
                dag, treatment=q.treatment, outcome=q.outcome, fmt="markdown",
                degradations=degradations,
            )
        except Exception as exc:
            record_degradation(
                degradations,
                section="causal DAG appendix (estimand-first path)",
                exc=exc,
                detail=f"dag_type={type(dag).__name__}",
            )

    # Build draft. workflow=None — this is the estimand-first path,
    # not the workflow path; provenance gets attached directly.
    draft = PaperDraft(
        question=eff_question,
        sections=sections,
        workflow=_LightweightWorkflowAdapter(q, result),
        fmt=fmt,
        citations=citations,
        parsed_hints={
            "treatment": q.treatment,
            "y": q.outcome,
            "design": q.design,
            "estimand": plan.estimand,
        },
        dag=dag,
        dag_treatment=q.treatment,
        dag_outcome=q.outcome,
        degradations=degradations,
    )

    # Attach provenance to the underlying estimator's result so
    # downstream replication_pack / Quarto appendix pick it up.
    try:
        from ..output._lineage import attach_provenance as _attach_prov
        target = result.underlying if result.underlying is not None else result
        _attach_prov(
            target,
            function=f"sp.causal_question[{plan.estimator}]",
            params={
                "treatment": q.treatment,
                "outcome": q.outcome,
                "estimand": plan.estimand,
                "design": q.design,
                "covariates": list(q.covariates) if q.covariates else None,
                "instruments": list(q.instruments) if q.instruments else None,
                "cutoff": q.cutoff,
                "time_structure": q.time_structure,
            },
            data=q.data,
            overwrite=False,
        )
    except Exception as exc:
        # Append directly to the already-built draft so introspection
        # works after construction.
        record_degradation(
            draft,
            section="provenance attachment (estimand-first path)",
            exc=exc,
            detail=f"estimator={plan.estimator}",
        )

    pipeline_notes = []
    for item in draft.degradations:
        detail = f" ({item['detail']})" if item.get("detail") else ""
        _record_note(
            pipeline_notes,
            f"{item['section']} degraded: {item['error_type']}: "
            f"{item['message']}{detail}",
        )
    if pipeline_notes:
        draft.sections["Pipeline notes"] = _notes_block(pipeline_notes)

    if output_path is not None:
        draft.write(output_path)

    return draft


class _LightweightWorkflowAdapter:
    """Minimal adapter so PaperDraft.to_qmd's _workflow_provenance() and
    replication_pack's _extract_data() / _extract_results() work for the
    estimand-first path (which has no real CausalWorkflow object).
    """

    __slots__ = ("question", "_estimation", "data", "result")

    def __init__(self, q, estimation):
        self.question = q
        self._estimation = estimation
        self.data = q.data
        # Surface the underlying estimator's result so the qmd
        # appendix and replication_pack pick up its _provenance.
        self.result = estimation.underlying if estimation.underlying is not None else estimation
