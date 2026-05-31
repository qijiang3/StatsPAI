"""
Modern IV reporting bundle — `sp.iv.iv_diag`.

This is the StatsPAI port of the post-2022 reporting standard for linear
IV in applied work, mirroring R `ivDiag` (Lal et al. 2024) and going a
step further by integrating the weak-IV-robust confidence sets
(:mod:`statspai.iv.weak_iv_ci`) and the Conley–Hansen–Rossi sensitivity
toolkit (:mod:`statspai.iv.plausibly_exogenous`) into a single object.

What it returns
---------------
:class:`IVDiagResult` — a structured bundle with:

- 2SLS (and optional OLS) point estimate with **analytic + bootstrap**
  standard errors (pairs and / or wild bootstrap; cluster-aware).
- Effective F (Olea–Pflueger 2013) and **tF-corrected critical value**
  (Lee–McCrary–Moreira–Porter 2022, AER 112, 3260–3290).
- Anderson–Rubin (1949) F at H0 and the AR confidence set; optional CLR
  / K confidence sets via :mod:`statspai.iv.weak_iv_ci`.
- Kleibergen–Paap (2006) rk LM and Wald F.
- Conley–Hansen–Rossi (2012) plausibly-exogenous LTZ sensitivity.
- A reading of the "TSLS-as-LATE" caveat
  (Blandhol–Bonney–Mogstad–Torgovitsky 2022/2025; Słoczyński 2024)
  when covariates are present and the endogenous regressor is binary.

The result object exposes ``.summary()``, ``.to_frame()``,
``.to_dict()``, ``.to_latex()``, ``.to_excel()``, ``.to_word()``, and
``.plot()`` for one-call manuscript-ready output.

References
----------
Anderson, T.W. and Rubin, H. (1949). [@anderson1949estimation]

Olea, J.L.M. and Pflueger, C. (2013). [@olea2013robust]

Conley, T.G., Hansen, C.B. and Rossi, P.E. (2012). [@conley2012plausibly]

Lee, D.S., McCrary, J., Moreira, M.J. and Porter, J. (2022).
    AER 112, 3260–3290. [@lee2022valid]

Young, A. (2022). EER 147, 104112. [@young2022consistency]

Keane, M.P. and Neal, T. (2024). Annual Review of Economics 16,
    185–212. [@keane2024practical]

Lal, A., Lockhart, M., Xu, Y. and Zu, Z. (2024). Political Analysis
    32(4), 521–540. [@lal2024much]

Blandhol, C., Bonney, J., Mogstad, M. and Torgovitsky, A. (2022/2025).
    NBER WP 29709. [@blandhol2025tsls]

Słoczyński, T. (2024). arXiv:2011.06695. [@sloczynski2024should]
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Sequence, Tuple, Union

import numpy as np
import pandas as pd

from ..diagnostics.weak_iv import (
    anderson_rubin_test,
    effective_f_test,
    tF_critical_value,
)


@dataclass
class IVDiagResult:
    """Container for :func:`iv_diag` output.

    Attributes
    ----------
    n : int
        Sample size after listwise deletion.
    n_endog, n_instruments, n_exog : int
        Counts of endogenous regressors, excluded instruments, included
        exogenous controls (excluding intercept).
    endog, instruments, exog : list[str]
        Variable names.
    beta_2sls, se_2sls, t_2sls, p_2sls : float
        2SLS point estimate, analytic SE, t-ratio, p-value (single
        endogenous regressor).
    ci_analytic_2sls : tuple[float, float]
        Analytic Wald CI based on ``se_2sls`` (level = ``1 - alpha``).
    beta_ols, se_ols, ci_ols : float, float, tuple[float, float]
        OLS counterpart (informative comparator; *not* causal).
    first_stage_F : float
        Classical first-stage F.
    effective_F : float
        Olea–Pflueger (2013) robust effective F.
    tF_critical_value : float
        Lee et al. (2022, AER) tF adjusted 5 % critical value at the
        observed first-stage F. ``inf`` if F < 3.84.
    ar_stat, ar_pvalue : float
        Anderson–Rubin (1949) F-statistic and p-value at ``h0``.
    ar_ci : tuple[float, float]
        AR confidence set (grid-inverted). ``±inf`` flags a one-/two-
        sided unbounded set.
    clr_ci, k_ci : tuple[float, float] | None
        Moreira (2003) CLR and Kleibergen (2002) K confidence sets.
        ``None`` if not requested.
    kp_rk_lm, kp_rk_lm_pvalue, kp_rk_f : float | None
        Kleibergen–Paap (2006) rk LM, p-value, Wald F.
    bootstrap_ci_analytic, bootstrap_ci_pairs, bootstrap_ci_wild
        Pair-/wild-bootstrap CI tuples (or ``None``).
    bootstrap_se_pairs, bootstrap_se_wild : float | None
        Bootstrap standard errors (matching CI sources).
    bootstrap_n : int
        Number of bootstrap replications actually used.
    ltz_ci, ltz_warning : tuple[float, float] | None, str | None
        Conley–Hansen–Rossi LTZ sensitivity CI under
        ``gamma_var = (gamma_sd) ** 2``.
    tF_adjusted_ci : tuple[float, float] | None
        ``beta ± tF_critical_value × se``. Falls back to ``±inf`` when
        ``F < 3.84`` (per LMMP 2022).
    tsls_late_caveat : str | None
        BBMT (2022/2025) / Słoczyński (2024) caveat text whenever the
        specification is at risk of negative-weight LATE pathologies.
    diagnostics : dict
        All numeric outputs in a flat dict (also returned by
        :meth:`to_dict`). Useful for downstream agent workflows.
    raw : dict
        Internal scratchpad with arrays (residuals, fitted values).
    """

    n: int
    n_endog: int
    n_instruments: int
    n_exog: int
    endog: str
    instruments: List[str]
    exog: List[str]
    alpha: float

    beta_2sls: float
    se_2sls: float
    t_2sls: float
    p_2sls: float
    ci_analytic_2sls: Tuple[float, float]

    beta_ols: float
    se_ols: float
    t_ols: float
    p_ols: float
    ci_ols: Tuple[float, float]

    first_stage_F: float
    effective_F: float
    tF_critical_value: float
    tF_adjusted_ci: Tuple[float, float]

    ar_stat: float
    ar_pvalue: float
    ar_ci: Tuple[float, float]

    kp_rk_lm: Optional[float]
    kp_rk_lm_pvalue: Optional[float]
    kp_rk_f: Optional[float]

    clr_ci: Optional[Tuple[float, float]] = None
    k_ci: Optional[Tuple[float, float]] = None

    bootstrap_n: int = 0
    bootstrap_se_pairs: Optional[float] = None
    bootstrap_ci_pairs: Optional[Tuple[float, float]] = None
    bootstrap_se_wild: Optional[float] = None
    bootstrap_ci_wild: Optional[Tuple[float, float]] = None

    ltz_ci: Optional[Tuple[float, float]] = None
    ltz_warning: Optional[str] = None

    tsls_late_caveat: Optional[str] = None

    diagnostics: Dict[str, Any] = field(default_factory=dict)
    raw: Dict[str, Any] = field(default_factory=dict)

    # ---------- summary table ---------------------------------------

    def to_frame(self) -> pd.DataFrame:
        """Return a tidy summary table — one row per estimator/metric."""
        rows = []
        ci_lo, ci_hi = self.ci_analytic_2sls
        rows.append(("2SLS (analytic)", self.beta_2sls, self.se_2sls,
                     self.t_2sls, self.p_2sls, ci_lo, ci_hi))
        if self.bootstrap_ci_pairs is not None:
            lo, hi = self.bootstrap_ci_pairs
            rows.append(("2SLS (pairs bootstrap)", self.beta_2sls,
                         self.bootstrap_se_pairs,
                         np.nan, np.nan, lo, hi))
        if self.bootstrap_ci_wild is not None:
            lo, hi = self.bootstrap_ci_wild
            rows.append(("2SLS (wild bootstrap)", self.beta_2sls,
                         self.bootstrap_se_wild,
                         np.nan, np.nan, lo, hi))
        lo, hi = self.tF_adjusted_ci
        rows.append(("2SLS (LMMP tF-adjusted)", self.beta_2sls, self.se_2sls,
                     np.nan, np.nan, lo, hi))
        lo, hi = self.ar_ci
        rows.append(("Anderson–Rubin set", np.nan, np.nan,
                     self.ar_stat, self.ar_pvalue, lo, hi))
        if self.clr_ci is not None:
            lo, hi = self.clr_ci
            rows.append(("Moreira CLR set", np.nan, np.nan,
                         np.nan, np.nan, lo, hi))
        if self.k_ci is not None:
            lo, hi = self.k_ci
            rows.append(("Kleibergen K set", np.nan, np.nan,
                         np.nan, np.nan, lo, hi))
        if self.ltz_ci is not None:
            lo, hi = self.ltz_ci
            rows.append(("CHR plausibly-exogenous LTZ", self.beta_2sls,
                         np.nan, np.nan, np.nan, lo, hi))
        ci_lo, ci_hi = self.ci_ols
        rows.append(("OLS (comparator, not causal)",
                     self.beta_ols, self.se_ols, self.t_ols,
                     self.p_ols, ci_lo, ci_hi))
        return pd.DataFrame(
            rows,
            columns=["estimator", "estimate", "SE", "stat", "p-value",
                     "CI lower", "CI upper"],
        )

    # ---------- summary text ----------------------------------------

    def summary(self) -> str:
        d = self
        lines: List[str] = []
        lines.append("Modern IV reporting bundle  (sp.iv.iv_diag)")
        lines.append("=" * 70)
        lines.append(f"  endog                       : {d.endog}")
        lines.append(f"  instruments                 : {', '.join(d.instruments)}")
        if d.exog:
            lines.append(f"  exogenous controls          : {', '.join(d.exog)}")
        lines.append(
            f"  N = {d.n}, k_endog = {d.n_endog}, k_instr = {d.n_instruments}, "
            f"k_exog = {d.n_exog}"
        )
        lines.append(f"  level                       : {100 * (1 - d.alpha):.0f} %")
        lines.append("-" * 70)
        lines.append("Strength:")
        lines.append(
            f"  First-stage F (classical)   : {d.first_stage_F:10.4f}"
        )
        lines.append(
            f"  Olea–Pflueger effective F   : {d.effective_F:10.4f}"
        )
        if d.kp_rk_f is not None:
            lines.append(
                f"  Kleibergen–Paap rk LM       : {d.kp_rk_lm:10.4f}"
                f"   p = {d.kp_rk_lm_pvalue:.4f}"
            )
            lines.append(
                f"  Kleibergen–Paap rk Wald F   : {d.kp_rk_f:10.4f}"
            )
        lines.append(
            f"  LMMP 2022 tF adjusted crit. : {d.tF_critical_value:10.4f}"
            f"   (F = {d.first_stage_F:.2f})"
        )
        lines.append("-" * 70)
        lines.append("Point estimate:")
        lines.append(
            f"  2SLS                        : {d.beta_2sls:10.4f}"
            f"   SE = {d.se_2sls:.4f}"
            f"   t = {d.t_2sls:.3f}   p = {d.p_2sls:.4f}"
        )
        ci_lo, ci_hi = d.ci_analytic_2sls
        lines.append(
            f"  Wald  {100 * (1 - d.alpha):.0f}% CI                : "
            f"[{ci_lo:.4f}, {ci_hi:.4f}]"
        )
        lo, hi = d.tF_adjusted_ci
        lines.append(
            f"  tF-corrected CI             : [{lo:.4f}, {hi:.4f}]"
        )
        if d.bootstrap_ci_pairs is not None:
            lo, hi = d.bootstrap_ci_pairs
            lines.append(
                f"  Pairs bootstrap (B={d.bootstrap_n})  : "
                f"SE = {d.bootstrap_se_pairs:.4f},  CI [{lo:.4f}, {hi:.4f}]"
            )
        if d.bootstrap_ci_wild is not None:
            lo, hi = d.bootstrap_ci_wild
            lines.append(
                f"  Wild bootstrap (B={d.bootstrap_n})    : "
                f"SE = {d.bootstrap_se_wild:.4f},  CI [{lo:.4f}, {hi:.4f}]"
            )
        lines.append("-" * 70)
        lines.append("Weak-IV-robust inference (size correct under any F):")
        lo, hi = d.ar_ci
        lines.append(
            f"  Anderson–Rubin F            : {d.ar_stat:10.4f}"
            f"   p = {d.ar_pvalue:.4f}"
        )
        lines.append(
            f"  AR  {100 * (1 - d.alpha):.0f}% confidence set    : "
            f"[{lo:.4f}, {hi:.4f}]"
        )
        if d.clr_ci is not None:
            lo, hi = d.clr_ci
            lines.append(
                f"  Moreira CLR  {100 * (1 - d.alpha):.0f}% set    : "
                f"[{lo:.4f}, {hi:.4f}]"
            )
        if d.k_ci is not None:
            lo, hi = d.k_ci
            lines.append(
                f"  Kleibergen K  {100 * (1 - d.alpha):.0f}% set   : "
                f"[{lo:.4f}, {hi:.4f}]"
            )
        if d.ltz_ci is not None:
            lo, hi = d.ltz_ci
            lines.append("-" * 70)
            lines.append("Plausibly exogenous (Conley–Hansen–Rossi 2012, LTZ):")
            lines.append(
                f"  LTZ  {100 * (1 - d.alpha):.0f}% CI            : "
                f"[{lo:.4f}, {hi:.4f}]"
            )
            if d.ltz_warning:
                lines.append(f"  Note                        : {d.ltz_warning}")
        lines.append("-" * 70)
        lines.append("OLS comparator (NOT a causal estimate):")
        ci_lo, ci_hi = d.ci_ols
        lines.append(
            f"  OLS                         : {d.beta_ols:10.4f}"
            f"   SE = {d.se_ols:.4f}   p = {d.p_ols:.4f}"
        )
        lines.append(
            f"  CI                          : [{ci_lo:.4f}, {ci_hi:.4f}]"
        )
        if d.tsls_late_caveat:
            lines.append("=" * 70)
            lines.append("⚠ Interpretation caveat:")
            lines.append("  " + d.tsls_late_caveat.replace("\n", "\n  "))
        lines.append("=" * 70)
        return "\n".join(lines)

    def __repr__(self) -> str:  # pragma: no cover
        return self.summary()

    # ---------- serialisation ---------------------------------------

    def to_dict(self) -> Dict[str, Any]:
        """Return all numeric diagnostics as a flat dict (jsonable)."""
        return dict(self.diagnostics)

    def to_latex(self, caption: Optional[str] = None,
                 label: Optional[str] = None,
                 float_format: str = "%.4f") -> str:
        """Render the summary table as a LaTeX ``tabular`` string."""
        df = self.to_frame()
        try:
            tex = df.to_latex(index=False, escape=False, na_rep="—",
                              float_format=float_format)
        except TypeError:
            # pandas < 1.4 fallback
            tex = df.to_latex(index=False, escape=False, na_rep="—")
        if caption or label:
            cap = f"\\caption{{{caption}}}\n" if caption else ""
            lab = f"\\label{{{label}}}\n" if label else ""
            tex = "\\begin{table}[!htbp]\n\\centering\n" + cap + lab + tex + "\\end{table}"
        return tex

    def to_excel(self, path: str) -> None:
        """Write the summary table to ``path`` (one sheet)."""
        self.to_frame().to_excel(path, index=False, sheet_name="iv_diag")

    def to_word(self, path: str, title: Optional[str] = None) -> None:
        """Write the summary table to a .docx file (requires python-docx)."""
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise ImportError(
                "to_word requires python-docx. Install with `pip install python-docx`."
            ) from exc
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        df = self.to_frame()
        table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
        table.style = "Light List"
        for j, col in enumerate(df.columns):
            table.rows[0].cells[j].text = str(col)
        for i, row in df.iterrows():
            for j, col in enumerate(df.columns):
                v = row[col]
                cell_text = (
                    "—" if (isinstance(v, float) and not np.isfinite(v))
                    else (f"{v:.4f}" if isinstance(v, float) else str(v))
                )
                table.rows[i + 1].cells[j].text = cell_text
        if self.tsls_late_caveat:
            doc.add_paragraph(
                "Caveat: " + self.tsls_late_caveat, style="Intense Quote"
            )
        doc.save(path)

    def plot(self, kind: str = "diagnostic", **kwargs):
        """Dispatch to :mod:`statspai.iv.plot` plotting helpers.

        Parameters
        ----------
        kind : {'diagnostic', 'forest', 'weak_iv', 'first_stage'}
            Which plot to render. ``'diagnostic'`` returns the 2x2 panel.
        """
        from . import plot as _ivplot  # lazy: matplotlib

        if kind == "diagnostic":
            return _ivplot.plot_iv_diagnostics(self, **kwargs)
        if kind == "forest":
            return _ivplot.plot_iv_forest_from_diag(self, **kwargs)
        if kind == "weak_iv":
            return _ivplot.plot_weak_iv_ci_overlay(self, **kwargs)
        if kind == "first_stage":
            raw = self.raw
            return _ivplot.plot_first_stage(
                endog=raw["D"], instruments=raw["Z"],
                exog=raw["W_no_const"], endog_name=self.endog,
                **kwargs,
            )
        raise ValueError(
            f"Unknown plot kind={kind!r}. Choose 'diagnostic', 'forest', "
            f"'weak_iv', or 'first_stage'."
        )


# ═══════════════════════════════════════════════════════════════════════
#  Internal helpers
# ═══════════════════════════════════════════════════════════════════════

def _prep_inputs(
    data: pd.DataFrame,
    y: str,
    endog: str,
    instruments: Sequence[str],
    exog: Optional[Sequence[str]] = None,
    cluster: Optional[Union[str, np.ndarray]] = None,
) -> Tuple[pd.DataFrame, np.ndarray, np.ndarray, np.ndarray, np.ndarray,
           Optional[np.ndarray]]:
    """Listwise-drop, return (df_clean, Y, D, Z, W_with_const, cluster_arr)."""
    instruments = list(instruments)
    exog_l = list(exog) if exog else []
    cluster_col = cluster if isinstance(cluster, str) else None
    cols = [y, endog] + instruments + exog_l
    if cluster_col is not None:
        cols = cols + [cluster_col]
    df_raw = data[cols].copy().dropna()
    # Save the surviving rows' positions (Int64Index) BEFORE we reset
    # the index — needed to align an external cluster array with the
    # post-dropna rows.
    _orig_idx = df_raw.index.to_numpy()
    df = df_raw.reset_index(drop=True)
    Y = df[y].to_numpy(dtype=float)
    D = df[endog].to_numpy(dtype=float)
    Z = df[instruments].to_numpy(dtype=float)
    if exog_l:
        Wx = df[exog_l].to_numpy(dtype=float)
        W = np.column_stack([np.ones(len(df)), Wx])
    else:
        W = np.ones((len(df), 1))
    if cluster_col is not None:
        cluster_arr = df[cluster_col].to_numpy()
    elif isinstance(cluster, (np.ndarray, pd.Series)):
        cluster_arr = np.asarray(cluster)
        if len(cluster_arr) != len(data):
            raise ValueError(
                f"cluster array length {len(cluster_arr)} does not match data "
                f"length {len(data)}."
            )
        # Align with the post-dropna rows by their *original* index
        # positions; reset_index(drop=True) above would otherwise alias
        # them to 0..n_clean-1 and silently shift cluster membership.
        cluster_arr = cluster_arr[_orig_idx]
    else:
        cluster_arr = None
    return df, Y, D, Z, W, cluster_arr


def _two_sls_point(Y: np.ndarray, D: np.ndarray, Z: np.ndarray,
                   W: np.ndarray) -> Tuple[float, np.ndarray, np.ndarray]:
    """Return (beta_endog, full coefficient vector, residuals)."""
    Z_full = np.column_stack([Z, W])
    X = np.column_stack([D.reshape(-1, 1), W])
    PZ_X, *_ = np.linalg.lstsq(Z_full, X, rcond=None)
    X_hat = Z_full @ PZ_X
    coef, *_ = np.linalg.lstsq(X_hat, Y, rcond=None)
    beta = float(coef[0])
    resid = Y - X @ coef  # 2SLS residuals computed at original X, not X_hat
    return beta, coef, resid


def _ols_point(Y: np.ndarray, D: np.ndarray, W: np.ndarray
               ) -> Tuple[float, float, np.ndarray]:
    """Return (beta_ols, se_ols, residuals)."""
    n = len(Y)
    X = np.column_stack([D.reshape(-1, 1), W])
    coef, *_ = np.linalg.lstsq(X, Y, rcond=None)
    beta = float(coef[0])
    resid = Y - X @ coef
    sigma2 = float(resid @ resid) / max(n - X.shape[1], 1)
    XtX_inv = np.linalg.pinv(X.T @ X)
    se = float(np.sqrt(sigma2 * XtX_inv[0, 0]))
    return beta, se, resid


def _se_2sls_robust(Y: np.ndarray, D: np.ndarray, Z: np.ndarray,
                    W: np.ndarray, beta: float,
                    cluster: Optional[np.ndarray] = None,
                    vcov: str = "HC1") -> float:
    """Heteroskedasticity-/cluster-robust 2SLS SE for the endog coefficient."""
    n = len(Y)
    Z_full = np.column_stack([Z, W])
    X = np.column_stack([D.reshape(-1, 1), W])
    PZ_X, *_ = np.linalg.lstsq(Z_full, X, rcond=None)
    X_hat = Z_full @ PZ_X
    XhX_inv = np.linalg.pinv(X_hat.T @ X)
    coef, *_ = np.linalg.lstsq(X_hat, Y, rcond=None)
    resid = Y - X @ coef
    # ── Classical homoskedastic 2SLS SE ───────────────────────────────
    if cluster is None and vcov == "classic":
        k = X.shape[1]
        sigma2 = float(resid @ resid) / max(n - k, 1)
        XhX_only_inv = np.linalg.pinv(X_hat.T @ X_hat)
        return float(np.sqrt(max(sigma2 * XhX_only_inv[0, 0], 0.0)))
    if cluster is None:
        if vcov == "HC0":
            scale = 1.0
        elif vcov == "HC1":
            scale = n / max(n - X.shape[1], 1)
        else:
            scale = 1.0
        meat = (X_hat * resid[:, None]).T @ (X_hat * resid[:, None]) * scale
    else:
        meat = np.zeros((X_hat.shape[1], X_hat.shape[1]))
        groups, _ = np.unique(cluster, return_inverse=True)
        G = len(groups)
        for g in groups:
            idx = (cluster == g)
            sg = (X_hat[idx] * resid[idx, None]).sum(axis=0)
            meat += np.outer(sg, sg)
        meat *= G / max(G - 1, 1) * (n - 1) / max(n - X.shape[1], 1)
    cov = XhX_inv @ meat @ XhX_inv.T
    return float(np.sqrt(max(cov[0, 0], 0.0)))


def _bootstrap_se(
    Y: np.ndarray,
    D: np.ndarray,
    Z: np.ndarray,
    W: np.ndarray,
    *,
    n_boot: int,
    cluster: Optional[np.ndarray],
    rng: np.random.Generator,
    method: str,
    alpha: float,
) -> Tuple[float, Tuple[float, float], int]:
    """Pairs (default) or wild bootstrap for the 2SLS coefficient.

    Returns (se_b, (lo, hi), B_used).
    """
    n = len(Y)
    betas = np.empty(n_boot)
    successes = 0
    if method == "pairs":
        if cluster is not None:
            groups = np.unique(cluster)
            for b in range(n_boot):
                draw = rng.choice(groups, size=len(groups), replace=True)
                idx = np.concatenate([np.where(cluster == g)[0] for g in draw])
                try:
                    bb, _, _ = _two_sls_point(Y[idx], D[idx], Z[idx], W[idx])
                except Exception:  # pragma: no cover
                    continue
                betas[successes] = bb
                successes += 1
        else:
            for b in range(n_boot):
                idx = rng.integers(0, n, size=n)
                try:
                    bb, _, _ = _two_sls_point(Y[idx], D[idx], Z[idx], W[idx])
                except Exception:  # pragma: no cover
                    continue
                betas[successes] = bb
                successes += 1
    elif method == "wild":
        # Davidson–MacKinnon (2010) wild bootstrap on 2SLS residuals.
        # y_star_i = X_i'β̂ + w_i · û_i, refit 2SLS.
        # Cluster-aware: one Rademacher weight per cluster.
        _, coef_hat, resid = _two_sls_point(Y, D, Z, W)
        X = np.column_stack([D.reshape(-1, 1), W])
        Yhat = X @ coef_hat
        for b in range(n_boot):
            if cluster is None:
                w = rng.choice([-1.0, 1.0], size=n)
            else:
                groups = np.unique(cluster)
                wg = rng.choice([-1.0, 1.0], size=len(groups))
                w = np.empty(n)
                for g_idx, g in enumerate(groups):
                    w[cluster == g] = wg[g_idx]
            Y_star = Yhat + w * resid
            try:
                bb, _, _ = _two_sls_point(Y_star, D, Z, W)
            except Exception:  # pragma: no cover
                continue
            betas[successes] = bb
            successes += 1
    else:
        raise ValueError(f"Unknown bootstrap method {method!r}.")

    if successes < max(50, n_boot // 4):
        return (float("nan"), (float("nan"), float("nan")), successes)
    sample = betas[:successes]
    se_b = float(np.std(sample, ddof=1))
    lo, hi = np.quantile(sample, [alpha / 2, 1 - alpha / 2])
    return (se_b, (float(lo), float(hi)), successes)


def _check_tsls_late_caveat(
    df: pd.DataFrame,
    endog: str,
    exog: Optional[Sequence[str]],
    instruments: Sequence[str],
) -> Optional[str]:
    """Return a caveat string when the spec is at risk of negative-weight LATE.

    Triggers when (a) covariates are present and (b) the endogenous regressor
    is binary 0/1 — the canonical Blandhol–Bonney–Mogstad–Torgovitsky and
    Słoczyński setup.
    """
    if not exog:
        return None
    d = df[endog].dropna()
    uniq = pd.unique(d)
    if len(uniq) > 2:
        return None
    if not set(np.unique(d.astype(float))).issubset({0.0, 1.0}):
        return None
    return (
        "TSLS with covariates and a binary endogenous regressor: linear 2SLS "
        "is in general NOT a positively-weighted average of conditional LATEs "
        "(Blandhol et al. 2022/2025; Słoczyński 2024). Negative weights on "
        "always-/never-takers are possible. To recover a properly weighted "
        "LATE / target-population effect, fit MTE bounds via "
        "`sp.iv(method='mte')` or `sp.iv(method='ivmte_bounds')`, or "
        "use a fully nonparametric (saturated) covariate specification."
    )


# ═══════════════════════════════════════════════════════════════════════
#  Public entry point
# ═══════════════════════════════════════════════════════════════════════

def iv_diag(
    data: pd.DataFrame,
    y: str,
    endog: str,
    instruments: Union[str, Sequence[str]],
    exog: Optional[Union[str, Sequence[str]]] = None,
    *,
    cluster: Optional[Union[str, np.ndarray]] = None,
    h0: float = 0.0,
    alpha: float = 0.05,
    vcov: str = "HC1",
    n_boot: int = 1000,
    boot_methods: Sequence[str] = ("pairs",),
    include_clr_ci: bool = False,
    include_k_ci: bool = False,
    grid_size: int = 401,
    ltz_gamma_sd: Optional[float] = None,
    random_state: Optional[int] = None,
) -> IVDiagResult:
    """
    Modern IV reporting bundle (single-endogenous, post-2022 standard).

    Returns a single :class:`IVDiagResult` containing:

    - 2SLS point estimate, analytic + bootstrap SEs, and Wald CI;
    - Olea–Pflueger effective F + Lee–McCrary–Moreira–Porter (2022) tF
      adjusted critical value and tF-corrected CI;
    - Anderson–Rubin (1949) F + AR confidence set; optional CLR / K
      confidence sets (Moreira 2003; Kleibergen 2002, 2005);
    - Kleibergen–Paap (2006) rk LM and Wald F;
    - Conley–Hansen–Rossi (2012) plausibly-exogenous LTZ sensitivity
      CI when ``ltz_gamma_sd`` is supplied;
    - the BBMT (2022/2025) / Słoczyński (2024) "TSLS-vs-LATE" caveat
      when covariates are present and the endogenous regressor is
      binary;
    - OLS comparator (informative; not causal).

    Parameters
    ----------
    data : DataFrame
    y : str
    endog : str
        Outcome and (single) endogenous regressor column names.
    instruments : str or list[str]
        Excluded instruments.
    exog : str or list[str], optional
        Included exogenous controls (intercept added automatically).
    cluster : str or array-like, optional
        Cluster variable for cluster-robust SEs and cluster bootstrap.
    h0 : float, default 0.0
        Null hypothesis value used by AR / CLR / K test.
    alpha : float, default 0.05
        Significance level for all CIs.
    vcov : {'HC0', 'HC1', 'classic'}
        Heteroskedasticity-robust covariance type for the analytic SE
        and the Olea–Pflueger effective F.
    n_boot : int, default 1000
        Bootstrap replications. Set to ``0`` to skip bootstrap.
    boot_methods : tuple of {'pairs', 'wild'}
        Which bootstraps to run.  Both can be requested.
    include_clr_ci, include_k_ci : bool
        Optionally invert CLR / K tests on a grid for the matching
        confidence set (slower; CLR uses Monte-Carlo critical values).
    grid_size : int
        Resolution of the AR / CLR / K grid inversion.
    ltz_gamma_sd : float, optional
        If supplied, run Conley–Hansen–Rossi LTZ sensitivity with prior
        ``γ ~ N(0, ltz_gamma_sd**2)``. Otherwise, LTZ is skipped.
    random_state : int, optional
        Seed for bootstrap and CLR Monte-Carlo.

    Returns
    -------
    IVDiagResult

    Examples
    --------
    >>> import statspai as sp
    >>> r = sp.iv.iv_diag(df, y='wage', endog='educ',
    ...                   instruments=['nearc4', 'nearc2'],
    ...                   exog=['exper', 'south'],
    ...                   n_boot=500, ltz_gamma_sd=0.05,
    ...                   random_state=42)
    >>> print(r.summary())
    >>> r.to_frame()                 # tidy table
    >>> r.plot('diagnostic')         # 2x2 diagnostic panel
    >>> r.to_latex(caption='IV diagnostic bundle')

    Notes
    -----
    - The bundle is **single-endogenous** by design; for multiple
      endogenous regressors use :func:`sp.iv.weakrobust` plus
      :func:`sp.iv.sanderson_windmeijer` per regressor.
    - The tF-corrected CI follows LMMP 2022: it widens (vs. the Wald
      CI) by exactly the multiplicative ratio
      ``c(F) / 1.96`` and equals ``±∞`` when ``F < 3.84``.
    - The bootstrap is implemented as a pairs (or wild Rademacher)
      bootstrap; see Young (2022, EER 147, 104112) for why analytic SEs
      can be unreliable. Cluster-bootstrap is used automatically when
      ``cluster`` is supplied.
    """
    if isinstance(instruments, str):
        instruments = [instruments]
    instruments = list(instruments)
    if exog is not None and isinstance(exog, str):
        exog = [exog]
    exog_l = list(exog) if exog else []

    df_clean, Y, D, Z, W, cluster_arr = _prep_inputs(
        data, y, endog, instruments, exog_l, cluster
    )
    n = len(df_clean)
    k_z = Z.shape[1]
    k_w = W.shape[1] - 1  # exclude constant

    # ── 2SLS point + analytic SE ──────────────────────────────────────
    beta_2sls, _coef_2sls, resid_2sls = _two_sls_point(Y, D, Z, W)
    se_2sls = _se_2sls_robust(Y, D, Z, W, beta_2sls,
                              cluster=cluster_arr, vcov=vcov)
    from scipy import stats
    t_2sls = beta_2sls / se_2sls if se_2sls > 0 else np.nan
    p_2sls = float(2 * (1 - stats.norm.cdf(abs(t_2sls)))) if np.isfinite(t_2sls) else np.nan
    z_crit = float(stats.norm.ppf(1 - alpha / 2))
    ci_analytic_2sls = (beta_2sls - z_crit * se_2sls,
                        beta_2sls + z_crit * se_2sls)

    # ── OLS comparator ────────────────────────────────────────────────
    beta_ols, se_ols, _ = _ols_point(Y, D, W)
    t_ols = beta_ols / se_ols if se_ols > 0 else np.nan
    p_ols = float(2 * (1 - stats.norm.cdf(abs(t_ols)))) if np.isfinite(t_ols) else np.nan
    ci_ols = (beta_ols - z_crit * se_ols, beta_ols + z_crit * se_ols)

    # ── AR / effective F / tF ─────────────────────────────────────────
    ar = anderson_rubin_test(
        data=df_clean, y=y, endog=endog,
        instruments=instruments, exog=exog_l or None,
        h0=h0, alpha=alpha, vcov=vcov,
    )
    first_stage_F = float(ar["first_stage_F"])
    effective_F = float(ar["effective_F"])
    # LMMP (2022) tF table is published only for alpha = 0.05 and F >= 3.84.
    # When either guard fails, the AR / weak-IV-robust set is the right
    # fallback and the t-ratio adjustment is meaningless — store inf.
    if not np.isclose(alpha, 0.05) or not np.isfinite(first_stage_F) \
            or first_stage_F < 3.84:
        tF_c = np.inf
        tF_ci = (-np.inf, np.inf)
    else:
        tF_c = tF_critical_value(first_stage_F, alpha=0.05)
        tF_ci = (beta_2sls - tF_c * se_2sls, beta_2sls + tF_c * se_2sls)

    # ── Kleibergen-Paap rk ────────────────────────────────────────────
    kp_rk_lm = kp_rk_lm_pvalue = kp_rk_f = None
    try:
        from .weak_identification import kleibergen_paap_rk
        kp = kleibergen_paap_rk(
            endog=D.reshape(-1, 1), instruments=Z,
            exog=W[:, 1:] if W.shape[1] > 1 else None,
            add_const=True,
            cov_type="cluster" if cluster_arr is not None else "robust",
            cluster=cluster_arr if cluster_arr is not None else None,
        )
        kp_rk_lm = float(kp.rk_lm)
        kp_rk_lm_pvalue = float(kp.rk_lm_pvalue)
        kp_rk_f = float(kp.rk_f)
    except Exception:  # pragma: no cover
        pass

    # ── Optional CLR / K confidence sets ──────────────────────────────
    clr_ci = k_ci = None
    if include_clr_ci or include_k_ci:
        try:
            from .weak_iv_ci import conditional_lr_ci, k_test_ci
            level = 1.0 - alpha
            if include_clr_ci:
                clr_cs = conditional_lr_ci(
                    y=y, endog=endog, instruments=instruments,
                    exog=exog_l or None, data=df_clean,
                    level=level, n_grid=grid_size,
                    n_sim=4000, random_state=random_state,
                )
                clr_ci = (float(clr_cs.lower), float(clr_cs.upper))
            if include_k_ci:
                k_cs = k_test_ci(
                    y=y, endog=endog, instruments=instruments,
                    exog=exog_l or None, data=df_clean,
                    level=level, n_grid=grid_size,
                )
                k_ci = (float(k_cs.lower), float(k_cs.upper))
        except Exception:  # pragma: no cover
            pass

    # ── Bootstrap ─────────────────────────────────────────────────────
    rng = np.random.default_rng(random_state)
    se_pairs = ci_pairs = se_wild = ci_wild = None
    boot_n_pairs = 0
    boot_n_wild = 0
    if n_boot > 0:
        for m in boot_methods:
            se_b, ci_b, used = _bootstrap_se(
                Y, D, Z, W, n_boot=n_boot, cluster=cluster_arr, rng=rng,
                method=m, alpha=alpha,
            )
            if m == "pairs":
                se_pairs, ci_pairs, boot_n_pairs = se_b, ci_b, used
            elif m == "wild":
                se_wild, ci_wild, boot_n_wild = se_b, ci_b, used
    # Headline `bootstrap_n`: if both methods ran, return min so a
    # conservative reader knows the weakest CI's replication count.
    if boot_n_pairs and boot_n_wild:
        boot_used = min(boot_n_pairs, boot_n_wild)
    else:
        boot_used = boot_n_pairs or boot_n_wild

    # ── Plausibly exogenous (LTZ) ─────────────────────────────────────
    ltz_ci = None
    ltz_warning = None
    if ltz_gamma_sd is not None:
        try:
            from .plausibly_exogenous import plausibly_exogenous_ltz
            ltz = plausibly_exogenous_ltz(
                y=y, endog=endog, instruments=instruments,
                exog=exog_l or None, data=df_clean,
                gamma_mean=0.0,
                gamma_var=float(ltz_gamma_sd) ** 2,
                ci_level=1.0 - alpha,
            )
            ltz_ci = (float(ltz.ci_lower), float(ltz.ci_upper))
            # CHR (2012) tipping-point: how large would |γ| need to be
            # in standard-deviation units of the prior before the LTZ
            # confidence set crosses 0?  Approximated by |β̂| / σ_γ.
            tip_sigmas = (abs(beta_2sls) / float(ltz_gamma_sd)
                          if float(ltz_gamma_sd) > 0 else float("inf"))
            ltz_warning = (
                f"Sensitivity prior γ ~ N(0, {float(ltz_gamma_sd):g}^2). "
                f"|β̂_2SLS| = {abs(beta_2sls):.4f} = "
                f"{tip_sigmas:.1f} × σ_γ — i.e. an exclusion-restriction "
                f"violation of that magnitude would zero out the IV estimate."
            )
        except Exception as exc:  # pragma: no cover
            ltz_warning = f"LTZ failed: {exc}"

    # ── TSLS-as-LATE caveat ───────────────────────────────────────────
    tsls_caveat = _check_tsls_late_caveat(df_clean, endog, exog_l, instruments)

    diagnostics: Dict[str, Any] = {
        "n": int(n),
        "n_endog": 1,
        "n_instruments": int(k_z),
        "n_exog": int(k_w),
        "alpha": float(alpha),
        "beta_2sls": float(beta_2sls),
        "se_2sls": float(se_2sls),
        "t_2sls": float(t_2sls) if np.isfinite(t_2sls) else None,
        "p_2sls": float(p_2sls) if np.isfinite(p_2sls) else None,
        "ci_analytic_2sls": [float(ci_analytic_2sls[0]), float(ci_analytic_2sls[1])],
        "beta_ols": float(beta_ols),
        "se_ols": float(se_ols),
        "ci_ols": [float(ci_ols[0]), float(ci_ols[1])],
        "first_stage_F": float(first_stage_F),
        "effective_F": float(effective_F),
        "tF_critical_value": float(tF_c),
        "tF_adjusted_ci": [float(tF_ci[0]), float(tF_ci[1])],
        "ar_stat": float(ar["ar_stat"]),
        "ar_pvalue": float(ar["ar_pvalue"]),
        "ar_ci": [float(ar["ar_ci"][0]), float(ar["ar_ci"][1])],
        "kp_rk_lm": kp_rk_lm,
        "kp_rk_lm_pvalue": kp_rk_lm_pvalue,
        "kp_rk_f": kp_rk_f,
        "clr_ci": list(clr_ci) if clr_ci else None,
        "k_ci": list(k_ci) if k_ci else None,
        "bootstrap_n": int(boot_used),
        "bootstrap_se_pairs": se_pairs,
        "bootstrap_ci_pairs": list(ci_pairs) if ci_pairs else None,
        "bootstrap_se_wild": se_wild,
        "bootstrap_ci_wild": list(ci_wild) if ci_wild else None,
        "ltz_ci": list(ltz_ci) if ltz_ci else None,
        "tsls_late_caveat": tsls_caveat,
    }

    return IVDiagResult(
        n=n, n_endog=1, n_instruments=k_z, n_exog=k_w,
        endog=endog, instruments=instruments, exog=exog_l,
        alpha=alpha,
        beta_2sls=beta_2sls, se_2sls=se_2sls,
        t_2sls=float(t_2sls) if np.isfinite(t_2sls) else float("nan"),
        p_2sls=float(p_2sls) if np.isfinite(p_2sls) else float("nan"),
        ci_analytic_2sls=ci_analytic_2sls,
        beta_ols=beta_ols, se_ols=se_ols,
        t_ols=float(t_ols) if np.isfinite(t_ols) else float("nan"),
        p_ols=float(p_ols) if np.isfinite(p_ols) else float("nan"),
        ci_ols=ci_ols,
        first_stage_F=first_stage_F,
        effective_F=effective_F,
        tF_critical_value=tF_c,
        tF_adjusted_ci=tF_ci,
        ar_stat=float(ar["ar_stat"]),
        ar_pvalue=float(ar["ar_pvalue"]),
        ar_ci=tuple(ar["ar_ci"]),
        kp_rk_lm=kp_rk_lm,
        kp_rk_lm_pvalue=kp_rk_lm_pvalue,
        kp_rk_f=kp_rk_f,
        clr_ci=clr_ci, k_ci=k_ci,
        bootstrap_n=boot_used,
        bootstrap_se_pairs=se_pairs,
        bootstrap_ci_pairs=ci_pairs,
        bootstrap_se_wild=se_wild,
        bootstrap_ci_wild=ci_wild,
        ltz_ci=ltz_ci, ltz_warning=ltz_warning,
        tsls_late_caveat=tsls_caveat,
        diagnostics=diagnostics,
        raw={
            "Y": Y, "D": D, "Z": Z, "W": W,
            "W_no_const": W[:, 1:] if W.shape[1] > 1 else None,
            "resid_2sls": resid_2sls,
            "cluster": cluster_arr,
        },
    )


# ═══════════════════════════════════════════════════════════════════════
#  Compare-across-methods convenience
# ═══════════════════════════════════════════════════════════════════════

def iv_compare(
    formula: Optional[str] = None,
    data: Optional[pd.DataFrame] = None,
    *,
    methods: Sequence[str] = ("2sls", "liml", "fuller", "jive"),
    alpha: float = 0.05,
    endog_name: Optional[str] = None,
    **kwargs,
) -> pd.DataFrame:
    """Run several k-class / JIVE estimators and return a one-row-per-method
    comparison table — useful for quick sensitivity checks across estimators.

    Parameters
    ----------
    formula : str
        IV formula (``"y ~ (d ~ z) + x"``).
    data : DataFrame
    methods : sequence of str
        Methods to dispatch through :func:`sp.iv` (and therefore the
        unified IV dispatcher).
    alpha : float
        For Wald-CI columns.

    Returns
    -------
    DataFrame
        columns ``method``, ``estimate``, ``SE``, ``CI lower``, ``CI upper``,
        ``first_stage_F``, ``effective_F``.
    """
    from . import _dispatch  # late import to avoid circulars
    from scipy import stats
    z_crit = float(stats.norm.ppf(1 - alpha / 2))

    # Pre-resolve endog name from the formula so we can look it up
    # consistently across heterogeneous result classes.
    canonical_endog = endog_name
    if canonical_endog is None and formula is not None and data is not None:
        try:
            from ..core.utils import parse_formula
            parsed = parse_formula(formula)
            ends = parsed.get("endogenous") or []
            if ends:
                canonical_endog = ends[0]
        except Exception:  # pragma: no cover
            pass

    rows = []
    for m in methods:
        try:
            res = _dispatch(formula=formula, data=data, method=m, **kwargs)
        except Exception as exc:  # pragma: no cover
            rows.append((m, np.nan, np.nan, np.nan, np.nan, np.nan, np.nan,
                         f"error: {exc}"))
            continue
        # Pull the endogenous-coefficient row
        try:
            params = res.params
            std_errors = res.std_errors
            diag = getattr(res, "diagnostics", {}) or {}
            local_endog = canonical_endog
            if local_endog is None or local_endog not in params.index:
                # Fall back: parse "First-stage F (<endog_name>)" diagnostic key.
                for key in diag.keys():
                    k_str = str(key)
                    if k_str.startswith("First-stage F (") and k_str.endswith(")"):
                        local_endog = k_str[len("First-stage F ("):-1]
                        if local_endog in params.index:
                            break
                        local_endog = None
            if local_endog is None and hasattr(res, "model") and \
                    hasattr(res.model, "_endog_names"):
                local_endog = res.model._endog_names[0]
            if local_endog is None:
                # last-resort: pick the LAST non-Intercept param name
                # (works for IVRegression; JIVEResult has endog FIRST,
                # but iv_compare really needs the canonical endog from
                # the formula — provide ``endog_name=`` to avoid ambiguity)
                for k in reversed(list(params.index)):
                    if "Intercept" not in str(k):
                        local_endog = k
                        break
            beta = float(params[local_endog])
            se = float(std_errors[local_endog])
            lo = beta - z_crit * se
            hi = beta + z_crit * se
            diag = getattr(res, "diagnostics", {}) or {}
            f_first = diag.get("First-stage F",
                               diag.get("first_stage_F", np.nan))
            f_eff = diag.get("Olea-Pflueger effective F",
                             diag.get("effective_F", np.nan))
            rows.append((m, beta, se, lo, hi,
                         float(f_first) if f_first not in (None, np.nan) else np.nan,
                         float(f_eff) if f_eff not in (None, np.nan) else np.nan,
                         "ok"))
        except Exception as exc:  # pragma: no cover
            rows.append((m, np.nan, np.nan, np.nan, np.nan, np.nan, np.nan,
                         f"summary error: {exc}"))
    return pd.DataFrame(
        rows,
        columns=["method", "estimate", "SE", "CI lower", "CI upper",
                 "first_stage_F", "effective_F", "status"],
    )


__all__ = ["iv_diag", "iv_compare", "IVDiagResult"]
