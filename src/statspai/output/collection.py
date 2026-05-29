"""``sp.Collection`` — session-level multi-table container.

Counterpart to Stata 15's ``collect`` and R's ``gt::gtsave`` —
gather any number of regression tables, descriptive statistics, balance
tables, and free-form text in one container, then export the whole
bundle to a single ``.docx`` / ``.xlsx`` / ``.tex`` / ``.md`` / ``.html``
file with consistent AER/QJE book-tab styling.

Examples
--------
>>> import statspai as sp
>>> c = sp.collect(title="Wage analysis", template="aer")
>>> c.add_regression(m1, m2, m3, name="main", title="Table 1: Wage equation")
>>> c.add_summary(df, vars=["wage", "educ", "exper"], title="Table 2: Descriptives")
>>> c.add_balance(df, treatment="treat", vars=["age", "female"],
...               title="Table 3: Balance")
>>> c.save("appendix.docx")
>>> c.save("appendix.xlsx")

The same collection produces a multi-section Word document, a workbook
with one sheet per item, and a single LaTeX / Markdown file that drops
straight into a manuscript.
"""

from __future__ import annotations

import warnings
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional, Sequence

import pandas as pd

from .regression_table import (
    MeanComparisonResult,
    RegtableResult,
    mean_comparison,
    regtable,
)
from .sumstats import sumstats

ItemKind = Literal["regtable", "summary", "balance", "tab", "text", "heading"]


@dataclass
class CollectionItem:
    """One entry in a :class:`Collection`."""

    name: str
    kind: ItemKind
    title: Optional[str]
    payload: Any  # RegtableResult / DataFrame / str / etc.
    options: Dict[str, Any] = field(default_factory=dict)


class Collection:
    """A named, ordered bundle of tables and prose for a single document.

    Add items in any order via the ``add_*`` methods (each returns
    ``self`` so calls can be chained). Render to any supported format
    via ``save(path)`` or one of the explicit ``to_*`` methods.

    Parameters
    ----------
    title : str, optional
        Displayed at the top of the rendered document.
    template : {'aer', 'qje', 'econometrica', 'restat'}, default 'aer'
        Forwarded to ``paper_tables`` style; also drives the default
        star levels used by ``add_regression``.
    """

    _NAME_RE = "[a-zA-Z0-9_\\-]+"

    def __init__(self, title: Optional[str] = None, template: str = "aer"):
        self.title = title
        self.template = template
        self.items: List[CollectionItem] = []

    # ------------------------------------------------------------------
    # Inspection
    # ------------------------------------------------------------------

    def __len__(self) -> int:
        return len(self.items)

    def __iter__(self):
        return iter(self.items)

    def __repr__(self) -> str:
        kinds = [it.kind for it in self.items]
        return (f"<Collection title={self.title!r} template={self.template!r} "
                f"items={len(self)} kinds={kinds}>")

    def list(self) -> pd.DataFrame:
        """Return a DataFrame summary (name / kind / title) for inspection."""
        rows = [{"name": it.name, "kind": it.kind, "title": it.title or ""}
                for it in self.items]
        return pd.DataFrame(rows, columns=["name", "kind", "title"])

    def to_frame(self, *, include_text: bool = False) -> pd.DataFrame:
        """Return a semantic long-format view of the collection.

        This is the programmatic counterpart to Stata's ``collect layout``
        and R's ``modelsummary``/``gt`` data pipeline: every rendered cell
        is represented as one row with stable dimensions
        ``item`` / ``kind`` / ``term`` / ``statistic`` / ``model`` and both
        a raw numeric ``value`` (when parseable) and the display
        ``formatted`` string.

        Parameters
        ----------
        include_text : bool, default False
            Include free-form text and headings as rows. Table workflows
            usually leave this off; document-audit workflows may want it.

        Returns
        -------
        pandas.DataFrame
            Long-format cell table with provenance-friendly dimensions.
        """
        rows: List[Dict[str, Any]] = []
        for item_index, it in enumerate(self.items):
            if it.kind in {"text", "heading"}:
                if include_text:
                    rows.append({
                        "collection": self.title or "",
                        "item": it.name,
                        "item_index": item_index,
                        "kind": it.kind,
                        "title": it.title or "",
                        "panel": it.options.get("panel", ""),
                        "model": "",
                        "term": "",
                        "statistic": "text",
                        "column": "",
                        "value": pd.NA,
                        "formatted": str(it.payload),
                        "source": type(it.payload).__name__,
                    })
                continue

            df = self._item_to_dataframe(it)
            if df is None or df.empty:
                continue
            for row_pos, (term, row) in enumerate(df.iterrows()):
                for col, cell in row.items():
                    statistic = str(col)
                    model = ""
                    if it.kind == "regtable":
                        model = str(col)
                        statistic = self._infer_regtable_statistic(str(term), str(cell))
                    rows.append({
                        "collection": self.title or "",
                        "item": it.name,
                        "item_index": item_index,
                        "kind": it.kind,
                        "title": it.title or "",
                        "panel": it.options.get("panel", ""),
                        "model": model,
                        "term": str(term),
                        "term_index": row_pos,
                        "statistic": statistic,
                        "column": str(col),
                        "value": self._coerce_numeric(cell),
                        "formatted": "" if pd.isna(cell) else str(cell),
                        "source": type(it.payload).__name__,
                    })
        columns = [
            "collection", "item", "item_index", "kind", "title", "panel",
            "model", "term", "term_index", "statistic", "column",
            "value", "formatted", "source",
        ]
        return pd.DataFrame(rows, columns=columns)

    # Alias mirroring tidyverse / broom vocabulary.
    to_long = to_frame

    def to_csv(self, path: Optional[str] = None, **kwargs) -> str:
        """Render :meth:`to_frame` to CSV; optionally write to ``path``."""
        content = self.to_frame().to_csv(index=False, **kwargs)
        if path is not None:
            Path(path).write_text(content, encoding="utf-8")
        return content

    @staticmethod
    def _serialize_payload(it: "CollectionItem") -> Dict[str, Any]:
        """JSON-safe serialisation of one item's payload.

        Regression tables reuse :meth:`RegtableResult.to_dict`; DataFrames
        (balance / summary / tab) round-trip through ``DataFrame.to_json``
        (NaN → null); text / heading carry their string.
        """
        import json
        p = it.payload
        if isinstance(p, RegtableResult):
            return p.to_dict()
        if isinstance(p, str):
            return {"text": p}
        if isinstance(p, pd.DataFrame):
            return {
                "columns": [str(c) for c in p.columns],
                "rows": json.loads(p.to_json(orient="records")),
            }
        # MeanComparisonResult / other result objects exposing a frame.
        if hasattr(p, "to_dataframe"):
            try:
                df = p.to_dataframe()
                return {
                    "columns": [str(c) for c in df.columns],
                    "rows": json.loads(df.to_json(orient="records")),
                }
            except Exception:
                pass
        return {"repr": str(p)}

    def to_dict(self) -> Dict[str, Any]:
        """Return a JSON-safe dict representation of the whole document.

        Agent-native counterpart to :meth:`save` — every item (regression
        table, balance / summary table, free text) is serialised in order so
        an LLM tool loop can cache and reason over a multi-table document
        without re-rendering. Regression-table items carry the full
        :meth:`RegtableResult.to_dict` payload (metadata + rendered grid +
        numeric truth).
        """
        return {
            "kind": "collection",
            "title": self.title,
            "template": self.template,
            "n_items": len(self.items),
            "items": [
                {
                    "name": it.name,
                    "item_kind": it.kind,
                    "title": it.title,
                    "content": self._serialize_payload(it),
                }
                for it in self.items
            ],
        }

    def to_json(self, *, indent: Optional[int] = None) -> str:
        """Serialise :meth:`to_dict` via ``json.dumps``."""
        import json
        return json.dumps(self.to_dict(), indent=indent, default=str)

    def get(self, name: str) -> CollectionItem:
        for it in self.items:
            if it.name == name:
                return it
        raise KeyError(f"No item named {name!r}; have {[i.name for i in self.items]}")

    def remove(self, name: str) -> "Collection":
        before = len(self.items)
        self.items = [it for it in self.items if it.name != name]
        if len(self.items) == before:
            raise KeyError(f"No item named {name!r}")
        return self

    def clear(self) -> "Collection":
        self.items.clear()
        return self

    @staticmethod
    def _coerce_numeric(value: Any) -> Any:
        if pd.isna(value):
            return pd.NA
        if isinstance(value, (int, float)):
            return float(value)
        s = str(value).strip()
        # Keep stars, parentheses, and thousands separators from blocking
        # numeric access to the underlying table cell.
        s = (
            s.replace(",", "")
            .replace("*", "")
            .replace("(", "")
            .replace(")", "")
        )
        try:
            return float(s)
        except ValueError:
            return pd.NA

    @staticmethod
    def _infer_regtable_statistic(term: str, formatted: str) -> str:
        t = term.lower()
        if t in {"n", "observations", "r-squared", "r2", "adj. r-squared"}:
            return t.replace(" ", "_").replace(".", "")
        if formatted.strip().startswith("(") and formatted.strip().endswith(")"):
            return "std_error"
        return "estimate"

    @staticmethod
    def _item_to_dataframe(it: CollectionItem) -> Optional[pd.DataFrame]:
        if it.kind == "summary":
            return it.payload.copy()
        payload = it.payload
        if hasattr(payload, "to_dataframe"):
            df = payload.to_dataframe()
            if isinstance(df, pd.DataFrame):
                return df.copy()
        if isinstance(payload, pd.DataFrame):
            return payload.copy()
        return None

    # ------------------------------------------------------------------
    # Adders
    # ------------------------------------------------------------------

    def _gen_name(self, prefix: str) -> str:
        used = {it.name for it in self.items}
        i = 1
        while f"{prefix}_{i}" in used:
            i += 1
        return f"{prefix}_{i}"

    def _check_name(self, name: str) -> None:
        if any(it.name == name for it in self.items):
            raise ValueError(f"item name {name!r} already in collection")

    def add_regression(
        self,
        *results,
        name: Optional[str] = None,
        title: Optional[str] = None,
        **regtable_kwargs,
    ) -> "Collection":
        """Add a regression table built from one or more model results.

        ``regtable_kwargs`` are forwarded verbatim to ``sp.regtable``.
        """
        if not results:
            raise ValueError("add_regression requires at least one model")
        rt = regtable(*results, title=title, **regtable_kwargs)
        nm = name or self._gen_name("regression")
        self._check_name(nm)
        self.items.append(CollectionItem(nm, "regtable", title, rt))
        return self

    def add_table(
        self,
        result,
        *,
        name: Optional[str] = None,
        title: Optional[str] = None,
    ) -> "Collection":
        """Add an already-built ``RegtableResult`` / ``MeanComparisonResult``."""
        if not isinstance(result, (RegtableResult, MeanComparisonResult)):
            raise TypeError(
                "add_table expects a RegtableResult or MeanComparisonResult; "
                f"got {type(result).__name__}"
            )
        kind: ItemKind = "regtable" if isinstance(result, RegtableResult) else "balance"
        nm = name or self._gen_name(kind)
        self._check_name(nm)
        if title and getattr(result, "title", None) in (None, ""):
            try:
                result.title = title
            except Exception:
                pass
        self.items.append(CollectionItem(nm, kind, title or getattr(result, "title", None), result))
        return self

    def add_summary(
        self,
        data: pd.DataFrame,
        vars: Optional[Sequence[str]] = None,
        *,
        stats: Optional[Sequence[str]] = None,
        name: Optional[str] = None,
        title: Optional[str] = None,
        labels: Optional[Dict[str, str]] = None,
    ) -> "Collection":
        """Add a descriptive-statistics table built from a DataFrame.

        Stores the underlying DataFrame; rendering re-uses the existing
        ``sumstats`` formatters so the AER book-tab style applies.
        """
        df = sumstats(
            data,
            vars=list(vars) if vars else None,
            stats=list(stats) if stats else None,
            labels=labels,
            output="dataframe",
            title=title,
        )
        nm = name or self._gen_name("summary")
        self._check_name(nm)
        self.items.append(CollectionItem(
            nm, "summary", title, df,
            options={"labels": labels or {}},
        ))
        return self

    def add_balance(
        self,
        data: pd.DataFrame,
        treatment: str,
        variables: Sequence[str],
        *,
        weights: Optional[str] = None,
        test: str = "ttest",
        name: Optional[str] = None,
        title: Optional[str] = None,
        fmt: str = "%.3f",
    ) -> "Collection":
        """Add a treatment vs. control balance table (calls ``mean_comparison``)."""
        result = mean_comparison(
            data,
            list(variables),
            group=treatment,
            weights=weights,
            test=test,
            fmt=fmt,
            title=title or "Balance",
        )
        nm = name or self._gen_name("balance")
        self._check_name(nm)
        self.items.append(CollectionItem(nm, "balance", title, result))
        return self

    def add_text(
        self,
        text: str,
        *,
        name: Optional[str] = None,
        title: Optional[str] = None,
    ) -> "Collection":
        """Add a free-form text block (rendered as a paragraph)."""
        if not isinstance(text, str):
            raise TypeError("add_text expects a str")
        nm = name or self._gen_name("text")
        self._check_name(nm)
        self.items.append(CollectionItem(nm, "text", title, text))
        return self

    def add_heading(
        self,
        text: str,
        *,
        level: int = 2,
        name: Optional[str] = None,
    ) -> "Collection":
        """Add a section heading (level 1-3)."""
        if not 1 <= level <= 3:
            raise ValueError("heading level must be 1, 2, or 3")
        nm = name or self._gen_name("heading")
        self._check_name(nm)
        self.items.append(CollectionItem(
            nm, "heading", text, text, options={"level": level}
        ))
        return self

    # ------------------------------------------------------------------
    # Renderers
    # ------------------------------------------------------------------

    def to_text(self) -> str:
        """Plain-text rendering of every item, top to bottom."""
        parts: List[str] = []
        if self.title:
            parts.append(self.title)
            parts.append("=" * len(self.title))
            parts.append("")
        for it in self.items:
            parts.append(self._item_to_text(it))
            parts.append("")
        return "\n".join(parts)

    def _item_to_text(self, it: CollectionItem) -> str:
        if it.kind == "heading":
            return f"\n{it.payload}\n{'-' * len(it.payload)}"
        if it.kind == "text":
            return it.payload
        if it.kind == "summary":
            head = f"=== {it.title or it.name} ===" if it.title or it.name else ""
            return f"{head}\n{it.payload.to_string()}".strip()
        if it.kind in {"regtable", "balance"}:
            return it.payload.to_text()
        return ""

    def to_markdown(self, path: Optional[str] = None) -> str:
        """Render to GitHub-flavoured Markdown; optionally write to ``path``."""
        parts: List[str] = []
        if self.title:
            parts.append(f"# {self.title}\n")
        for it in self.items:
            if it.kind == "heading":
                lvl = it.options.get("level", 2)
                parts.append(f"{'#' * lvl} {it.payload}\n")
            elif it.kind == "text":
                if it.title:
                    parts.append(f"## {it.title}\n")
                parts.append(f"{it.payload}\n")
            elif it.kind == "summary":
                if it.title:
                    parts.append(f"## {it.title}\n")
                parts.append(it.payload.to_markdown() + "\n")
            elif it.kind in {"regtable", "balance"}:
                if it.title:
                    parts.append(f"## {it.title}\n")
                parts.append(it.payload.to_markdown() + "\n")
        content = "\n".join(parts)
        if path is not None:
            Path(path).write_text(content, encoding="utf-8")
        return content

    def to_html(self, path: Optional[str] = None) -> str:
        """Render to a single self-contained HTML document."""
        parts: List[str] = ["<html><head><meta charset='utf-8'>"]
        if self.title:
            parts.append(f"<title>{self.title}</title>")
        parts.append("<style>body{font-family:'Times New Roman',serif;}"
                     "table{border-collapse:collapse;margin:1em 0;}"
                     "th,td{padding:4px 10px;}"
                     "table{border-top:2px solid #000;border-bottom:2px solid #000;}"
                     "th{border-bottom:1px solid #000;}</style></head><body>")
        if self.title:
            parts.append(f"<h1>{self.title}</h1>")
        for it in self.items:
            if it.kind == "heading":
                lvl = it.options.get("level", 2)
                parts.append(f"<h{lvl}>{it.payload}</h{lvl}>")
            elif it.kind == "text":
                if it.title:
                    parts.append(f"<h2>{it.title}</h2>")
                parts.append(f"<p>{it.payload}</p>")
            elif it.kind == "summary":
                if it.title:
                    parts.append(f"<h2>{it.title}</h2>")
                parts.append(it.payload.to_html())
            elif it.kind in {"regtable", "balance"}:
                if it.title:
                    parts.append(f"<h2>{it.title}</h2>")
                parts.append(it.payload.to_html())
        parts.append("</body></html>")
        content = "\n".join(parts)
        if path is not None:
            Path(path).write_text(content, encoding="utf-8")
        return content

    def to_latex(self, path: Optional[str] = None) -> str:
        """Concatenate every item's LaTeX into one ``.tex`` file."""
        parts: List[str] = []
        if self.title:
            parts.append(f"% {self.title}")
        for it in self.items:
            if it.kind == "heading":
                lvl = it.options.get("level", 2)
                section = {1: "section", 2: "subsection", 3: "subsubsection"}[lvl]
                parts.append(f"\\{section}{{{it.payload}}}\n")
            elif it.kind == "text":
                if it.title:
                    parts.append(f"\\subsection*{{{it.title}}}")
                parts.append(it.payload + "\n")
            elif it.kind == "summary":
                if it.title:
                    parts.append(f"\\subsection*{{{it.title}}}")
                parts.append(it.payload.to_latex() + "\n")
            elif it.kind in {"regtable", "balance"}:
                parts.append(it.payload.to_latex() + "\n")
        content = "\n".join(parts)
        if path is not None:
            Path(path).write_text(content, encoding="utf-8")
        return content

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def to_docx(self, path: str) -> str:
        """Write the entire collection to a single ``.docx`` file.

        Each item renders in turn — headings as Word headings, text as
        paragraphs, tables in AER book-tab style — separated by page
        breaks between tables.
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as e:
            raise ImportError(
                "python-docx is required for .docx export. "
                "Install with: pip install python-docx"
            ) from e

        from ._aer_style import (
            apply_word_booktab_rules,
            style_word_table_typography,
            add_word_notes_paragraph,
        )

        doc = Document()
        if self.title:
            t = doc.add_heading(self.title, level=1)
            for run in t.runs:
                run.font.name = "Times New Roman"

        last_was_table = False
        for it in self.items:
            if it.kind == "heading":
                lvl = it.options.get("level", 2)
                h = doc.add_heading(it.payload, level=lvl)
                for run in h.runs:
                    run.font.name = "Times New Roman"
                last_was_table = False
                continue
            if it.kind == "text":
                if it.title:
                    h = doc.add_heading(it.title, level=2)
                    for run in h.runs:
                        run.font.name = "Times New Roman"
                p = doc.add_paragraph()
                run = p.add_run(it.payload)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                last_was_table = False
                continue

            # Tables: heading + DataFrame -> AER booktab table
            if it.title:
                h = doc.add_heading(it.title, level=2)
                for run in h.runs:
                    run.font.name = "Times New Roman"

            if it.kind == "summary":
                df = it.payload.copy()
                # sumstats puts the variable name in the index already
                df_to_render = df.reset_index()
                df_to_render.columns = ["Variable"] + list(df.columns)
                self._write_aer_table_to_doc(
                    doc, df_to_render,
                    apply_word_booktab_rules, style_word_table_typography,
                    notes=None,
                )
            elif it.kind == "balance":
                df = it.payload.to_dataframe().reset_index()
                self._write_aer_table_to_doc(
                    doc, df,
                    apply_word_booktab_rules, style_word_table_typography,
                    notes="* p<0.10, ** p<0.05, *** p<0.01",
                    add_notes=add_word_notes_paragraph,
                )
            else:  # regtable
                rt = it.payload
                df = rt.to_dataframe().reset_index()
                # `to_dataframe()` uses unnamed index → reset_index gives 'index' col.
                # Rename to '' so the first column has no header.
                first_col = df.columns[0]
                df = df.rename(columns={first_col: ""})
                note_lines: List[str] = []
                try:
                    note_lines.append(f"{rt._se_label()} in parentheses")
                except Exception:
                    pass
                if getattr(rt, "show_stars", True):
                    try:
                        note_lines.append(rt._star_note())
                    except Exception:
                        note_lines.append("* p<0.10, ** p<0.05, *** p<0.01")
                for n in getattr(rt, "notes", []) or []:
                    note_lines.append(n)
                self._write_aer_table_to_doc(
                    doc, df,
                    apply_word_booktab_rules, style_word_table_typography,
                    notes="\n".join(note_lines),
                    add_notes=add_word_notes_paragraph,
                )
            last_was_table = True

        doc.save(path)
        return path

    @staticmethod
    def _write_aer_table_to_doc(
        doc, df, apply_rules, style_typo,
        *, notes: Optional[str] = None, add_notes=None,
    ) -> None:
        n_rows = len(df) + 1
        n_cols = len(df.columns)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.autofit = True
        for j, col in enumerate(df.columns):
            table.rows[0].cells[j].text = str(col)
        for i, (_, row_data) in enumerate(df.iterrows(), 1):
            for j, val in enumerate(row_data):
                table.rows[i].cells[j].text = "" if pd.isna(val) else str(val)
        style_typo(table, header_rows=(0,))
        apply_rules(table, header_top_idx=0, header_bot_idx=0)
        if notes and add_notes is not None:
            add_notes(doc, notes)

    # ------------------------------------------------------------------
    # XLSX
    # ------------------------------------------------------------------

    def to_xlsx(self, path: str) -> str:
        """Write the collection to a single workbook (one sheet per item)."""
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment
        except ImportError as e:
            raise ImportError(
                "openpyxl is required for .xlsx export. "
                "Install with: pip install openpyxl"
            ) from e

        from ._aer_style import excel_booktab_borders

        top_rule, mid_rule, bottom_rule, _ = excel_booktab_borders()
        header_font = Font(bold=True, name="Times New Roman", size=11)
        body_font = Font(name="Times New Roman", size=11)
        title_font = Font(bold=True, name="Times New Roman", size=12)
        notes_font = Font(italic=True, name="Times New Roman", size=9)
        center = Alignment(horizontal="center")

        wb = openpyxl.Workbook()
        wb.remove(wb.active)

        if self.title:
            ws = wb.create_sheet(title="overview")
            ws.cell(row=1, column=1, value=self.title).font = title_font

        for it in self.items:
            if it.kind == "heading":
                continue  # headings collapse into the next sheet's title
            sheet_name = it.name[:31]
            ws = wb.create_sheet(title=sheet_name)
            row = 1
            if it.title:
                ws.cell(row=row, column=1, value=str(it.title)).font = title_font
                row += 2
            if it.kind == "text":
                ws.cell(row=row, column=1, value=str(it.payload)).font = body_font
                continue

            if it.kind == "summary":
                df = it.payload.reset_index()
                df.columns = ["Variable"] + list(df.columns[1:])
            elif it.kind == "balance":
                df = it.payload.to_dataframe().reset_index()
            else:  # regtable
                df = it.payload.to_dataframe().reset_index()
                df.columns = [""] + list(df.columns[1:])

            self._write_aer_table_to_sheet(
                ws, df, row,
                top_rule=top_rule, mid_rule=mid_rule, bottom_rule=bottom_rule,
                header_font=header_font, body_font=body_font, center=center,
            )

            # Notes
            note_row = row + len(df) + 2
            if it.kind in {"regtable", "balance"}:
                payload = it.payload
                try:
                    se_label = payload._se_label()
                    ws.cell(row=note_row, column=1,
                            value=f"{se_label} in parentheses").font = notes_font
                    note_row += 1
                except Exception:
                    pass
                if getattr(payload, "show_stars", True):
                    try:
                        ws.cell(row=note_row, column=1,
                                value=payload._star_note()).font = notes_font
                        note_row += 1
                    except Exception:
                        pass
                for n in getattr(payload, "notes", []) or []:
                    ws.cell(row=note_row, column=1, value=str(n)).font = notes_font
                    note_row += 1

            for col_cells in ws.columns:
                width = max((len(str(c.value)) for c in col_cells if c.value), default=8)
                ws.column_dimensions[col_cells[0].column_letter].width = min(width + 3, 28)

        wb.save(path)
        return path

    @staticmethod
    def _write_aer_table_to_sheet(
        ws, df, start_row, *, top_rule, mid_rule, bottom_rule,
        header_font, body_font, center,
    ) -> None:
        # Header
        for j, col in enumerate(df.columns, 1):
            cell = ws.cell(row=start_row, column=j, value=str(col))
            cell.font = header_font
            cell.alignment = center
            cell.border = top_rule
        # Mid rule
        for j in range(1, len(df.columns) + 1):
            ws.cell(row=start_row + 1, column=j).border = mid_rule
        # Body
        for i, (_, row_data) in enumerate(df.iterrows()):
            r = start_row + 2 + i
            for j, val in enumerate(row_data, 1):
                cell = ws.cell(row=r, column=j, value="" if pd.isna(val) else str(val))
                cell.font = body_font
                if j > 1:
                    cell.alignment = center
        # Bottom rule
        last_row = start_row + 1 + len(df)
        for j in range(1, len(df.columns) + 1):
            ws.cell(row=last_row, column=j).border = bottom_rule

    # ------------------------------------------------------------------
    # save (auto-detect)
    # ------------------------------------------------------------------

    def save(self, path: str) -> str:
        """Auto-detect format from ``path`` extension and write."""
        ext = Path(path).suffix.lower()
        if ext == ".docx":
            return self.to_docx(path)
        if ext in (".xlsx", ".xls"):
            return self.to_xlsx(path)
        if ext in (".tex", ".latex"):
            self.to_latex(path)
            return path
        if ext in (".html", ".htm"):
            self.to_html(path)
            return path
        if ext in (".md", ".markdown"):
            self.to_markdown(path)
            return path
        if ext == ".txt":
            Path(path).write_text(self.to_text(), encoding="utf-8")
            return path
        raise ValueError(
            f"Unsupported extension {ext!r}. "
            "Choose from .docx, .xlsx, .tex, .html, .md, .txt."
        )


def collect(title: Optional[str] = None, *, template: str = "aer") -> Collection:
    """Construct a fresh :class:`Collection`.

    Convenience factory mirroring Stata 15's ``collect`` workflow:

    >>> import statspai as sp
    >>> c = sp.collect("Wage analysis")
    >>> c.add_regression(m1, m2, name="main")
    >>> c.add_summary(df, vars=["wage", "educ"], name="desc")
    >>> c.save("paper.docx")
    """
    return Collection(title=title, template=template)
