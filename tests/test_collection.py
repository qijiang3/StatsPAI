"""
Tests for ``sp.Collection`` / ``sp.collect`` — the session-level
multi-table container.

Mirrors the Stata 15 ``collect`` / R ``gt::gtsave`` workflow: gather
heterogeneous results into one container, export to a single file.
"""

import numpy as np
import pandas as pd
import pytest

import statspai as sp


@pytest.fixture
def models_and_df():
    rng = np.random.default_rng(33)
    n = 400
    treat = rng.integers(0, 2, n)
    x1 = rng.normal(size=n)
    x2 = rng.normal(size=n)
    y = 1.0 + 0.7 * treat + 0.4 * x1 - 0.2 * x2 + rng.normal(0, 1, n)
    df = pd.DataFrame({"y": y, "treat": treat, "x1": x1, "x2": x2})
    m1 = sp.regress("y ~ treat", data=df)
    m2 = sp.regress("y ~ treat + x1", data=df)
    m3 = sp.regress("y ~ treat + x1 + x2", data=df)
    return df, m1, m2, m3


# ---------------------------------------------------------------------------
# Construction & inspection
# ---------------------------------------------------------------------------


def test_collect_factory_returns_collection():
    c = sp.collect("Wage analysis", template="aer")
    assert isinstance(c, sp.Collection)
    assert c.title == "Wage analysis"
    assert c.template == "aer"
    assert len(c) == 0


def test_collection_chained_adds_yield_self(models_and_df):
    df, m1, m2, _ = models_and_df
    c = sp.collect()
    out = c.add_regression(m1, m2, name="main").add_summary(
        df, vars=["x1", "x2"], name="desc"
    )
    assert out is c
    assert len(c) == 2
    assert [it.name for it in c] == ["main", "desc"]


def test_collection_list_returns_dataframe(models_and_df):
    df, m1, m2, _ = models_and_df
    c = sp.collect()
    c.add_regression(m1, m2, name="reg")
    c.add_summary(df, vars=["x1"], name="sum")
    c.add_text("Notes here.", name="note")
    c.add_heading("Section A", name="hd")
    listing = c.list()
    assert set(listing.columns) == {"name", "kind", "title"}
    assert set(listing["name"]) == {"reg", "sum", "note", "hd"}
    assert set(listing["kind"]) == {"regtable", "summary", "text", "heading"}


def test_collection_remove_by_name(models_and_df):
    _, m1, _, _ = models_and_df
    c = sp.collect().add_regression(m1, name="main").add_text("hi", name="note")
    c.remove("note")
    assert [it.name for it in c] == ["main"]
    with pytest.raises(KeyError):
        c.remove("does-not-exist")


def test_collection_duplicate_name_raises(models_and_df):
    _, m1, m2, _ = models_and_df
    c = sp.collect().add_regression(m1, name="main")
    with pytest.raises(ValueError, match="already in collection"):
        c.add_regression(m2, name="main")


def test_collection_get_lookup(models_and_df):
    _, m1, _, _ = models_and_df
    c = sp.collect().add_regression(m1, name="main")
    assert c.get("main").kind == "regtable"
    with pytest.raises(KeyError):
        c.get("nope")


def test_collection_to_frame_exposes_semantic_cells(models_and_df):
    df, m1, m2, _ = models_and_df
    c = (
        sp.collect("Doc")
        .add_regression(m1, m2, name="main", title="Main")
        .add_summary(df, vars=["x1"], name="desc")
        .add_text("Audit note", name="note")
    )
    long = c.to_frame()
    assert {
        "item", "kind", "model", "term", "statistic", "value", "formatted",
    } <= set(long.columns)
    assert {"main", "desc"} <= set(long["item"])
    assert "note" not in set(long["item"])
    assert long.query("item == 'main'")["model"].astype(str).str.len().gt(0).any()
    assert pd.to_numeric(long["value"], errors="coerce").notna().any()


def test_collection_to_frame_can_include_text(models_and_df):
    _, m1, _, _ = models_and_df
    c = sp.collect().add_heading("Appendix").add_regression(m1, name="main")
    long = c.to_frame(include_text=True)
    assert "heading" in set(long["kind"])
    assert "Appendix" in set(long["formatted"])


def test_collection_to_csv_writes_long_table(tmp_path, models_and_df):
    _, m1, _, _ = models_and_df
    c = sp.collect().add_regression(m1, name="main")
    out = tmp_path / "collect.csv"
    csv_text = c.to_csv(str(out))
    assert out.exists()
    assert "item, item_index" not in csv_text  # no accidental spaces in header
    assert "formatted" in csv_text


# ---------------------------------------------------------------------------
# Adders
# ---------------------------------------------------------------------------


def test_add_regression_requires_at_least_one_model(models_and_df):
    c = sp.collect()
    with pytest.raises(ValueError, match="at least one model"):
        c.add_regression(name="empty")


def test_add_balance_calls_mean_comparison(models_and_df):
    df, *_ = models_and_df
    c = sp.collect().add_balance(df, treatment="treat", variables=["x1", "x2"], name="bal")
    item = c.get("bal")
    assert item.kind == "balance"


def test_add_summary_stores_dataframe(models_and_df):
    df, *_ = models_and_df
    c = sp.collect().add_summary(df, vars=["x1", "x2"], name="desc")
    item = c.get("desc")
    assert item.kind == "summary"
    assert isinstance(item.payload, pd.DataFrame)


def test_add_heading_validates_level(models_and_df):
    c = sp.collect()
    with pytest.raises(ValueError):
        c.add_heading("bad", level=4)


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------


def test_to_text_includes_all_items(models_and_df):
    df, m1, m2, _ = models_and_df
    c = sp.collect("Doc")
    c.add_heading("Part 1")
    c.add_regression(m1, m2, name="main", title="Table 1: Main")
    c.add_summary(df, vars=["x1", "x2"], name="desc", title="Table 2: Descriptives")
    c.add_text("Standard errors clustered at the firm level.", name="note")
    text = c.to_text()
    assert "Doc" in text
    assert "Table 1" in text
    assert "Table 2" in text
    assert "clustered at the firm level" in text


def test_to_markdown_writes_file(tmp_path, models_and_df):
    _, m1, m2, _ = models_and_df
    c = sp.collect().add_regression(m1, m2, name="main", title="Main")
    out = tmp_path / "doc.md"
    c.to_markdown(str(out))
    text = out.read_text(encoding="utf-8")
    assert "Main" in text


def test_to_html_self_contained(tmp_path, models_and_df):
    _, m1, m2, _ = models_and_df
    out = tmp_path / "doc.html"
    sp.collect("Doc").add_regression(m1, m2, name="main").to_html(str(out))
    html = out.read_text(encoding="utf-8")
    assert html.startswith("<html>")
    assert "<title>Doc</title>" in html


def test_save_auto_detects_extension(tmp_path, models_and_df):
    _, m1, m2, _ = models_and_df
    c = sp.collect().add_regression(m1, m2, name="main")
    md = tmp_path / "out.md"
    tex = tmp_path / "out.tex"
    c.save(str(md))
    c.save(str(tex))
    assert md.exists() and md.stat().st_size > 50
    assert tex.exists() and tex.stat().st_size > 50
    # Both renderings must reference the regressor we used
    assert "treat" in md.read_text(encoding="utf-8")
    assert "treat" in tex.read_text(encoding="utf-8")


def test_save_unsupported_ext_raises(models_and_df, tmp_path):
    _, m1, _, _ = models_and_df
    c = sp.collect().add_regression(m1, name="main")
    with pytest.raises(ValueError, match="Unsupported"):
        c.save(str(tmp_path / "x.unknown"))


# ---------------------------------------------------------------------------
# DOCX / XLSX
# ---------------------------------------------------------------------------


def test_save_docx(tmp_path, models_and_df):
    docx = pytest.importorskip("docx")
    df, m1, m2, m3 = models_and_df

    c = (
        sp.collect("Wage analysis")
        .add_heading("Section A: Main")
        .add_regression(m1, m2, m3, name="main", title="Table 1")
        .add_summary(df, vars=["x1", "x2"], name="desc", title="Table 2")
        .add_balance(df, treatment="treat", variables=["x1", "x2"],
                     name="bal", title="Table 3")
        .add_text("Standard errors clustered at firm level.", name="note")
    )
    out = tmp_path / "everything.docx"
    c.save(str(out))
    assert out.exists() and out.stat().st_size > 1000

    from docx import Document
    doc = Document(str(out))
    assert len(doc.tables) >= 3  # main + summary + balance
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Wage analysis" in text
    assert "clustered at firm level" in text


def test_save_xlsx_one_sheet_per_item(tmp_path, models_and_df):
    openpyxl = pytest.importorskip("openpyxl")
    df, m1, m2, _ = models_and_df
    c = (
        sp.collect("Doc")
        .add_regression(m1, m2, name="main")
        .add_summary(df, vars=["x1", "x2"], name="desc")
    )
    out = tmp_path / "doc.xlsx"
    c.save(str(out))
    wb = openpyxl.load_workbook(str(out))
    sheets = set(wb.sheetnames)
    assert "main" in sheets
    assert "desc" in sheets


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------


def test_collect_is_in_registry():
    funcs = sp.list_functions()
    names = {f["name"] if isinstance(f, dict) else f for f in funcs}
    assert "collect" in names
