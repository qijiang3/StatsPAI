"""Coverage-focused tests for decomposition internals.

Targets currently-uncovered lines in
``statspai.decomposition._common`` and ``statspai.decomposition._results``.
Every test carries a real assertion (closed-form identity, algebraic
invariant, exact structural property, or a triggered ``raise``) — no bare
smoke calls, no mocked numerical paths.
"""
from __future__ import annotations

import warnings

import numpy as np
import pandas as pd
import pytest

import statspai as sp  # noqa: F401  (import-alias contract)

from statspai.decomposition._common import (
    add_constant,
    wls,
    cluster_vcov,
    bootstrap_stat,
    wild_bootstrap_stat,
    bootstrap_ci,
    weighted_quantile,
    weighted_ecdf,
    kde_at,
    weighted_gini,
    statistic_value,
    influence_function,
    prepare_frame,
)
from statspai.decomposition._results import (
    DecompResultMixin,
    _coerce_for_json,
)

SEED = 20240607


# ════════════════════════════════════════════════════════════════════════
# _common.py — cluster_vcov  (lines 80-94)
# ════════════════════════════════════════════════════════════════════════

def test_cluster_vcov_single_cluster_per_obs_matches_hc1_shape():
    """cluster_vcov is symmetric PSD-shaped and uses the CR1 finite-sample
    factor. With every observation in its OWN cluster (G == n) the meat is
    the sum of outer products of per-obs scores — a well-defined matrix."""
    rng = np.random.default_rng(SEED)
    n, k = 60, 3
    X = add_constant(rng.normal(size=(n, k - 1)))
    beta_true = np.array([1.0, -2.0, 0.5])
    y = X @ beta_true + rng.normal(scale=0.7, size=n)
    beta, _, resid = wls(y, X, robust=True)

    clusters = np.arange(n)  # each obs its own cluster
    V = cluster_vcov(X, resid, clusters)

    assert V.shape == (k, k)
    # symmetric
    assert np.allclose(V, V.T, atol=1e-10)
    # diagonal variances strictly positive
    assert np.all(np.diag(V) > 0)


def test_cluster_vcov_weights_default_and_grouping():
    """Two genuine clusters; default weights path (w is None -> ones).
    Reordering observations within clusters must not change V (the meat is
    a within-cluster sum, order-invariant)."""
    rng = np.random.default_rng(SEED + 1)
    n, k = 80, 2
    X = add_constant(rng.normal(size=(n, k - 1)))
    y = X @ np.array([0.5, 1.5]) + rng.normal(size=n)
    beta, _, resid = wls(y, X, robust=False)
    clusters = np.repeat([0, 1, 2, 3], n // 4)

    V1 = cluster_vcov(X, resid, clusters)  # w None branch

    perm = rng.permutation(n)
    V2 = cluster_vcov(X[perm], resid[perm], clusters[perm])
    assert np.allclose(V1, V2, atol=1e-9)
    assert np.allclose(V1, V1.T, atol=1e-12)


# ════════════════════════════════════════════════════════════════════════
# _common.py — bootstrap_stat  (lines 213, 234-236, 238, 241)
# ════════════════════════════════════════════════════════════════════════

def test_bootstrap_stat_default_rng_is_deterministic():
    """rng=None -> default_rng(12345); calling twice yields identical draws
    (line 213)."""
    out1 = bootstrap_stat(lambda idx: float(idx.mean()), n=50, n_boot=30)
    out2 = bootstrap_stat(lambda idx: float(idx.mean()), n=50, n_boot=30)
    assert out1.shape == (30, 1)
    assert np.array_equal(out1, out2)  # deterministic seed


def test_bootstrap_stat_partial_failures_warn_and_recover():
    """Some replications raise -> counted, skipped, and a RuntimeWarning is
    emitted because >5% failed (lines 234-236, 241). Survivors still yield
    a usable array."""
    rng = np.random.default_rng(SEED)
    calls = {"i": 0}

    def flaky(idx):
        calls["i"] += 1
        if calls["i"] % 2 == 0:  # ~50% fail -> exceeds 5% threshold
            raise ValueError("boom")
        return float(idx.mean())

    with pytest.warns(RuntimeWarning, match="bootstrap replications failed"):
        out = bootstrap_stat(flaky, n=40, n_boot=40, rng=rng)
    # roughly half survived
    assert 1 <= out.shape[0] < 40
    assert out.shape[1] == 1


def test_bootstrap_stat_all_fail_raises():
    """Every replication raises -> RuntimeError (lines 237-238)."""
    def always_fail(idx):
        raise RuntimeError("nope")

    with pytest.raises(RuntimeError, match="All bootstrap replications failed"):
        bootstrap_stat(always_fail, n=10, n_boot=5,
                       rng=np.random.default_rng(0))


# ════════════════════════════════════════════════════════════════════════
# _common.py — wild_bootstrap_stat  (lines 289, 315, 320-322, 324, 326)
# ════════════════════════════════════════════════════════════════════════

def test_wild_bootstrap_default_rng_and_rademacher_symmetry():
    """rng=None default (line 289). With Rademacher multipliers and a
    symmetric stat_fn (sum of pseudo-residual deviations from fitted), the
    mean over replications stays near the baseline statistic."""
    n = 40
    fitted = np.linspace(0, 1, n)
    resid = np.full(n, 0.3)

    def stat(y_star):
        return float((y_star - fitted).sum())  # == sum(v_i * resid_i)

    out = wild_bootstrap_stat(stat, resid, fitted, n_boot=200,
                              weights="rademacher")
    assert out.shape == (200, 1)
    # E[v]=0 => mean of sum(v*resid) ~ 0; bounded by |sum(resid)|
    assert abs(out.mean()) < resid.sum()


def test_wild_bootstrap_unknown_weights_raises():
    """Unknown multiplier scheme -> ValueError (line 315)."""
    n = 10
    with pytest.raises(ValueError, match="unknown weights"):
        wild_bootstrap_stat(
            lambda y: float(y.sum()),
            resid=np.ones(n), fitted=np.zeros(n),
            n_boot=3, weights="gaussian",
            rng=np.random.default_rng(0),
        )


def test_wild_bootstrap_partial_failures_warn():
    """>5% replications raise -> counted + RuntimeWarning (320-322, 326)."""
    calls = {"i": 0}

    def flaky(y_star):
        calls["i"] += 1
        if calls["i"] % 2 == 0:
            raise ValueError("x")
        return float(y_star.sum())

    with pytest.warns(RuntimeWarning, match="wild-bootstrap replications failed"):
        out = wild_bootstrap_stat(
            flaky, resid=np.ones(20), fitted=np.zeros(20),
            n_boot=40, weights="rademacher",
            rng=np.random.default_rng(SEED),
        )
    assert out.shape[1] == 1
    assert out.shape[0] >= 1


def test_wild_bootstrap_all_fail_raises():
    """Every replication raises -> RuntimeError (lines 323-324)."""
    with pytest.raises(RuntimeError, match="All wild-bootstrap replications failed"):
        wild_bootstrap_stat(
            lambda y: (_ for _ in ()).throw(ValueError("bad")),
            resid=np.ones(8), fitted=np.zeros(8),
            n_boot=4, weights="rademacher",
            rng=np.random.default_rng(0),
        )


# ════════════════════════════════════════════════════════════════════════
# _common.py — bootstrap_ci  (line 362) + methods
# ════════════════════════════════════════════════════════════════════════

def test_bootstrap_ci_transpose_branch_for_vector_point():
    """A (n_boot, 1)-shaped boot with a multi-element point triggers the
    transpose reshape (line 362): boot (B,1) -> (1,B). With a single row per
    column the quantiles collapse onto the single value, giving a degenerate
    but well-formed CI (lo == hi == that value)."""
    B = 5
    boot = np.arange(B, dtype=float).reshape(B, 1)  # (5,1)
    point = np.array([10.0, 20.0, 30.0])            # size 3 > 1
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", RuntimeWarning)  # ddof on single row
        se, lo, hi = bootstrap_ci(boot, point, method="percentile")
    # After transpose boot is (1, B): one obs per column.
    assert se.shape == (B,)
    assert lo.shape == (B,)
    # percentile of a single value is that value
    assert np.allclose(lo, np.arange(B))
    assert np.allclose(hi, np.arange(B))


def test_bootstrap_ci_basic_and_normal_methods():
    """'basic' reflects quantiles around 2*point; 'normal' is point ± z*se;
    unknown method raises."""
    rng = np.random.default_rng(SEED)
    boot = rng.normal(loc=2.0, scale=1.0, size=(500, 1))
    point = np.array([2.0])

    se_b, lo_b, hi_b = bootstrap_ci(boot, point, method="basic")
    se_n, lo_n, hi_n = bootstrap_ci(boot, point, method="normal")
    # normal CI is exactly symmetric about point
    assert np.allclose((lo_n + hi_n) / 2, point, atol=1e-9)
    assert lo_b[0] < hi_b[0]

    with pytest.raises(ValueError, match="unknown method"):
        bootstrap_ci(boot, point, method="bca")


# ════════════════════════════════════════════════════════════════════════
# _common.py — weighted_quantile / weighted_ecdf / kde_at  (403, 413, 434)
# ════════════════════════════════════════════════════════════════════════

def test_weighted_quantile_scalar_returns_float():
    """Scalar q -> python float (line 403). The cumulative-weight CDF runs
    from w/Σw up to 1, so np.interp at q=0.5 over 1..9 lands at 4.5; the
    point of the test is the scalar->float coercion plus monotonicity."""
    y = np.arange(1, 10, dtype=float)
    med = weighted_quantile(y, 0.5)
    assert isinstance(med, float)
    assert med == pytest.approx(4.5)
    # vector q stays an array (no scalar branch)
    qs = weighted_quantile(y, np.array([0.1, 0.9]))
    assert isinstance(qs, np.ndarray)
    assert qs[0] < qs[1]


def test_weighted_ecdf_default_weights_monotone():
    """w None -> unit weights (line 413). ECDF is non-decreasing and hits
    1.0 at the max of the sample."""
    y_sample = np.array([1.0, 2.0, 3.0, 4.0])
    grid = np.array([0.5, 1.0, 2.5, 4.0, 10.0])
    F = weighted_ecdf(grid, y_sample)
    assert F[0] == 0.0                     # below min
    assert np.all(np.diff(F) >= -1e-12)    # monotone non-decreasing
    assert F[-1] == pytest.approx(1.0)     # at/above max


def test_kde_at_degenerate_sigma_fallback():
    """Constant data => weighted cov is 0 => sigma fallback path (line 434).
    Density at the constant point stays finite and positive."""
    y = np.full(20, 3.0)
    d = kde_at(y, 3.0)
    assert np.isfinite(d)
    assert d > 0


# ════════════════════════════════════════════════════════════════════════
# _common.py — weighted_gini  (lines 452, 457)
# ════════════════════════════════════════════════════════════════════════

def test_weighted_gini_zero_total_weight_nan():
    """W <= 0 -> nan (line 452)."""
    y = np.array([1.0, 2.0, 3.0])
    w = np.zeros(3)
    assert np.isnan(weighted_gini(y, w))


def test_weighted_gini_nonpositive_mean_nan():
    """mu <= 0 -> nan (line 457). Symmetric-about-zero data has mean 0."""
    y = np.array([-2.0, -1.0, 1.0, 2.0])
    w = np.ones(4)
    assert np.isnan(weighted_gini(y, w))


# ════════════════════════════════════════════════════════════════════════
# _common.py — statistic_value  (lines 495, 501, 507) + branches
# ════════════════════════════════════════════════════════════════════════

@pytest.mark.parametrize("stat", ["theil_t", "theil_l", "atkinson"])
def test_statistic_value_inequality_positive_data(stat):
    """theil_t / theil_l / atkinson clip y to 1e-12 before averaging, so the
    plug-in mean is always positive and the ``mu <= 0`` nan guards
    (lines 495/501/507) are unreachable through the public surface — see
    report. Here we exercise the positive-data body and assert each index is
    a finite, non-negative float (all three inequality measures are >= 0)."""
    rng = np.random.default_rng(SEED)
    y = rng.lognormal(size=200)
    w = np.ones_like(y)
    val = statistic_value(y, w, stat)
    assert np.isfinite(val)
    assert isinstance(val, float)
    assert val >= -1e-9


def test_statistic_value_full_menu_and_unknown_raises():
    """Exercise the remaining statistic_value branches with closed-form
    identities, and the unknown-stat ValueError."""
    y = np.array([1.0, 2.0, 3.0, 4.0, 5.0])
    w = np.ones_like(y)
    assert statistic_value(y, w, "mean") == pytest.approx(3.0)
    # ddof=1 variance of 1..5 == 2.5
    assert statistic_value(y, w, "variance") == pytest.approx(2.5)
    assert statistic_value(y, w, "std") == pytest.approx(np.sqrt(2.5))
    # quantile delegates to weighted_quantile (CDF-inversion convention)
    assert statistic_value(y, w, "quantile", tau=0.5) == pytest.approx(
        float(weighted_quantile(y, 0.5, w=w))
    )
    assert statistic_value(y, w, "iqr") == pytest.approx(
        float(weighted_quantile(y, 0.75, w=w) - weighted_quantile(y, 0.25, w=w))
    )
    assert np.isfinite(statistic_value(y, w, "gini"))
    assert np.isfinite(statistic_value(y, w, "log_var"))
    with pytest.raises(ValueError, match="unknown statistic"):
        statistic_value(y, w, "median_absolute_deviation")


# ════════════════════════════════════════════════════════════════════════
# _common.py — influence_function  (608, 620, 628, 634, 641)
# ════════════════════════════════════════════════════════════════════════

def test_influence_function_gini_nonpositive_mean_returns_nan():
    """gini does NOT clip y, so mean-zero (symmetric) data gives mu == 0
    and the IF is all-nan (line 608). The theil_t/theil_l/atkinson guards
    (620/628/634) clip y to 1e-12 first, so their mu<=0 paths are
    unreachable through the public surface — see report."""
    y = np.array([-3.0, -1.0, 1.0, 3.0])
    rif = influence_function(y, "gini")
    assert rif.shape == y.shape
    assert np.all(np.isnan(rif))


@pytest.mark.parametrize("stat", ["gini", "theil_t", "theil_l", "atkinson"])
def test_influence_function_positive_data_recenters_to_statistic(stat):
    """Recentering property: the mean of a RIF equals the plug-in statistic.
    This exercises the full positive-data body of each inequality IF
    (gini lines through 615, theil_t through 623, theil_l 629, atkinson
    635-640)."""
    rng = np.random.default_rng(SEED)
    y = rng.lognormal(mean=0.0, sigma=0.4, size=4000)
    w = np.ones_like(y)
    rif = influence_function(y, stat, w=w)
    assert rif.mean() == pytest.approx(statistic_value(y, w, stat), rel=0.02)


def test_influence_function_mean_identity_and_unknown_raises():
    """IF of the mean is y itself; unknown stat raises (line 641)."""
    y = np.array([2.0, 4.0, 6.0])
    assert np.array_equal(influence_function(y, "mean"), y)
    with pytest.raises(ValueError, match="unknown statistic"):
        influence_function(y, "kurtosis")


def test_influence_function_quantile_recenters_to_quantile():
    """RIF of a quantile averages back to the quantile (recentering
    property): E[RIF] ≈ q."""
    rng = np.random.default_rng(SEED)
    y = rng.normal(size=2000)
    rif = influence_function(y, "quantile", tau=0.5)
    q = weighted_quantile(y, 0.5)
    assert rif.mean() == pytest.approx(q, abs=0.05)


# ════════════════════════════════════════════════════════════════════════
# _common.py — prepare_frame  (line 679)
# ════════════════════════════════════════════════════════════════════════

def test_prepare_frame_weight_length_mismatch_raises():
    """An array of weights whose length differs from the cleaned frame
    raises (line 679)."""
    df = pd.DataFrame({"y": [1.0, 2.0, 3.0], "x": [0.1, 0.2, 0.3]})
    bad_w = np.ones(2)  # wrong length
    with pytest.raises(ValueError, match="weights array length"):
        prepare_frame(df, ["y", "x"], weights=bad_w)


def test_prepare_frame_string_weights_and_dropna():
    """String weights column is extracted and dropped; NA rows removed."""
    df = pd.DataFrame({
        "y": [1.0, 2.0, np.nan, 4.0],
        "x": [0.1, 0.2, 0.3, 0.4],
        "wt": [1.0, 1.0, 1.0, 2.0],
    })
    out, w = prepare_frame(df, ["y", "x"], weights="wt")
    assert "wt" not in out.columns
    assert len(out) == 3                # NA row dropped
    assert np.array_equal(w, np.array([1.0, 1.0, 2.0]))


# ════════════════════════════════════════════════════════════════════════
# _results.py — DecompResultMixin via a tiny concrete subclass
# ════════════════════════════════════════════════════════════════════════

class _Res(DecompResultMixin):
    method_name = "ToyDecomp"
    bib_keys = ("oaxaca1973male",)

    def __init__(self, **kw):
        for k, v in kw.items():
            setattr(self, k, v)


def test_confint_detailed_branch_full():
    """which='detailed' with a 'se' column and a 'contribution' value column
    builds per-variable CIs (covers the detailed branch incl. lines 337-351;
    exercises the value-col detection)."""
    detailed = pd.DataFrame({
        "variable": ["age", "edu"],
        "contribution": [0.5, -0.2],
        "se": [0.1, 0.05],
    })
    r = _Res(detailed=detailed)
    ci = r.confint(which="detailed")
    assert set(ci) == {"age", "edu"}
    lo, hi = ci["age"]
    z = 1.959963984540054
    assert lo == pytest.approx(0.5 - z * 0.1, abs=1e-6)
    assert hi == pytest.approx(0.5 + z * 0.1, abs=1e-6)


def test_confint_detailed_empty_or_no_se_returns_none():
    """Empty detailed -> None (line 334); detailed without 'se' -> None
    (line 336)."""
    assert _Res(detailed=pd.DataFrame()).confint(which="detailed") is None
    no_se = pd.DataFrame({"variable": ["a"], "contribution": [1.0]})
    assert _Res(detailed=no_se).confint(which="detailed") is None


def test_confint_detailed_no_value_column_returns_none():
    """A 'se' column but no recognised value column -> None (line 344)."""
    df = pd.DataFrame({"variable": ["a"], "se": [0.1], "other": [1.0]})
    assert _Res(detailed=df).confint(which="detailed") is None


def test_confint_unknown_which_raises():
    r = _Res(overall={"gap": 1.0, "gap_se": 0.2})
    with pytest.raises(ValueError, match="unknown which"):
        r.confint(which="sideways")


def test_cite_no_keys_returns_empty():
    """No bib_keys -> '' for string fmt, [] otherwise (line 368)."""
    class _NoKeys(DecompResultMixin):
        bib_keys = ()
    nk = _NoKeys()
    assert nk.cite("string") == ""
    assert nk.cite("list") == []
    assert nk.cite("bibtex_keys") == []


def test_cite_resolves_known_key():
    r = _Res()
    s = r.cite("string")
    assert "Oaxaca" in s and "1973" in s
    assert r.cite("bibtex_keys") == ["oaxaca1973male"]


def test_to_excel_no_panels_raises():
    """A result with no exportable panels -> RuntimeError (line 426)."""
    empty = _Res()  # no overall, no detailed, no scalars
    with pytest.raises(RuntimeError, match="no exportable panels"):
        empty.to_excel()


def test_to_excel_roundtrip_bytes():
    """to_excel(path=None) returns workbook bytes; sheets are readable."""
    r = _Res(
        overall={"gap": 1.0, "explained": 0.6, "unexplained": 0.4},
        detailed=pd.DataFrame({"variable": ["x"], "contribution": [0.6]}),
    )
    data = r.to_excel()
    assert isinstance(data, (bytes, bytearray))
    import io
    book = pd.read_excel(io.BytesIO(data), sheet_name=None)
    assert "Overall" in book
    assert "Detailed" in book


def test_dataframe_panels_scalar_attrs(tmp_path):
    """Canonical scalar attributes (no 'overall' dict) populate the Overall
    panel from scalar_rows (lines 544, 546). Confirm via to_excel sheet."""
    r = _Res(observed_gap=2.0, closed_gap=0.5, total_change=1.5)
    panels = r._dataframe_panels()
    assert "Overall" in panels
    quantities = set(panels["Overall"]["quantity"])
    assert {"observed_gap", "closed_gap", "total_change"} <= quantities


def test_to_word_empty_dataframe_panel(tmp_path):
    """A DataFrame panel that is empty after detection... empty panels are
    filtered by _dataframe_panels, so we cover the empty-table branch
    (lines 487-488) by overriding _dataframe_panels to yield an empty df."""
    class _WordRes(DecompResultMixin):
        method_name = "EmptyPanelDoc"
        bib_keys = ("oaxaca1973male",)

        def _dataframe_panels(self):
            return {"Empty": pd.DataFrame(columns=["a", "b"]),
                    "Full": pd.DataFrame({"a": [1], "b": [2.5]})}

    out = tmp_path / "doc.docx"
    res = _WordRes().to_word(str(out))
    assert out.exists()
    from docx import Document
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "(empty)" in texts  # empty-panel paragraph (line 487-488)


def test_to_word_summary_exception_falls_back_to_repr(tmp_path):
    """If summary() raises, to_word swallows it and writes repr(self)
    instead (lines 476-477)."""
    class _BadSummary(DecompResultMixin):
        method_name = "BadSummaryDoc"
        bib_keys = ()

        def _dataframe_panels(self):
            return {}

        def summary(self):
            raise RuntimeError("summary blew up")

        def __repr__(self):
            return "REPR_FALLBACK_MARKER"

    out = tmp_path / "bad.docx"
    _BadSummary().to_word(str(out))
    from docx import Document
    doc = Document(str(out))
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "REPR_FALLBACK_MARKER" in full


def test_to_word_writes_summary_and_tables(tmp_path):
    """to_word writes a heading, the summary text, a populated table, and a
    References section. summary() returns text (no exception => normal
    path)."""
    class _SumRes(DecompResultMixin):
        method_name = "SummaryDoc"
        bib_keys = ("oaxaca1973male",)
        detailed = pd.DataFrame({"variable": ["x"], "contribution": [1.25]})

        def summary(self):
            return "Gap = 1.25\nExplained = 1.00"

    out = tmp_path / "sum.docx"
    _SumRes().to_word(str(out))
    from docx import Document
    doc = Document(str(out))
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Gap = 1.25" in full
    assert "References" in full
    assert any(t.rows for t in doc.tables)


# ════════════════════════════════════════════════════════════════════════
# _results.py — _coerce_for_json  (lines 575, 579)
# ════════════════════════════════════════════════════════════════════════

def test_coerce_for_json_ndarray_to_list():
    """numpy ndarray -> python list (line 575)."""
    out = _coerce_for_json({"a": np.array([1.0, 2.0, 3.0])})
    assert out == {"a": [1.0, 2.0, 3.0]}
    assert isinstance(out["a"], list)


def test_coerce_for_json_nan_and_inf_to_none():
    """Bare python float nan/inf -> None (line 579)."""
    assert _coerce_for_json(float("nan")) is None
    assert _coerce_for_json(float("inf")) is None
    assert _coerce_for_json(2.5) == 2.5


def test_to_dict_roundtrips_through_json():
    """to_dict / to_json coerce numpy + DataFrame into JSON-safe forms."""
    import json
    r = _Res(
        overall={"gap": np.float64(1.0)},
        detailed=pd.DataFrame({"variable": ["x"], "contribution": [0.5]}),
        arr=np.array([1, 2, 3]),
    )
    d = r.to_dict()
    assert d["arr"] == [1, 2, 3]
    parsed = json.loads(r.to_json())
    assert parsed["overall"]["gap"] == 1.0
